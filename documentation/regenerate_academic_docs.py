from __future__ import annotations

import json
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "documentation"
DOC_DIR.mkdir(parents=True, exist_ok=True)
UPLOADS_DIR = ROOT / "uploads"
PAPERS_FILE = ROOT / "docs" / "paper_summary.md"

PROJECT_TITLE = "NAS Medical Model Generator"
FORMAL_SUBTITLE = "Neural Architecture Search Based Medical Prediction System"


def new_document(title: str, subtitle: str) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(18)

    subtitle_paragraph = doc.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_paragraph.add_run(subtitle)
    run.italic = True
    run.font.size = Pt(14)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"Generated on: {date.today().isoformat()}")
    doc.add_page_break()
    return doc


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)


def add_long_paragraphs(doc: Document, paragraphs: Iterable[str]) -> None:
    for paragraph in paragraphs:
        add_paragraph(doc, paragraph)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_toc(doc: Document, items: Iterable[tuple[str, int, str]]) -> None:
    add_heading(doc, "TABLE OF CONTENTS", 1)
    for title, indent, page in items:
        paragraph = doc.add_paragraph()
        if indent:
            paragraph.paragraph_format.left_indent = Pt(indent * 18)
        paragraph.add_run(f"{title} ........................................ {page}")


def add_report_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row_values in rows:
        row = table.add_row().cells
        for index, value in enumerate(row_values):
            row[index].text = str(value)


def metric_percent(value: object) -> str:
    if isinstance(value, (int, float)):
        return f"{value * 100:.2f}%"
    return str(value)


def format_architecture(arch: object) -> str:
    if isinstance(arch, dict):
        dense_units = arch.get("dense_units", [])
        activation = arch.get("activation", "unknown")
        if dense_units:
            return f"{len(dense_units)} hidden layers with units {dense_units} and {activation} activation"
    return str(arch)


def load_training_report() -> dict:
    report_path = UPLOADS_DIR / "training_report.json"
    if report_path.exists():
        return json.loads(report_path.read_text(encoding="utf-8"))
    return {"task_type": "classification", "all_models": [], "best_model": {}}


def load_paper_lines(limit: int = 16) -> list[dict]:
    if not PAPERS_FILE.exists():
        return []

    lines = PAPERS_FILE.read_text(encoding="utf-8").splitlines()
    papers: list[dict] = []
    current: dict | None = None
    current_field: str | None = None
    field_map = {
        "authors": "authors",
        "year": "year",
        "source": "source",
        "pdf": "pdf",
        "abstract": "abstract",
        "advantages": "advantages",
        "datasets": "datasets",
        "disadvantages": "disadvantages",
        "results": "results",
        "relevance to our project": "relevance",
        "relevance": "relevance",
    }

    for raw_line in lines:
        line = raw_line.strip()
        if line.startswith("## "):
            if current:
                papers.append(current)
            current = {"title": line[3:].strip()}
            current_field = None
            continue
        if not current:
            continue
        if line.startswith("- ") and ":" in line:
            key, value = line[2:].split(":", 1)
            normalized = field_map.get(key.strip().lower())
            if normalized:
                current[normalized] = value.strip()
                current_field = normalized
            continue
        if current_field and line:
            current[current_field] = f"{current.get(current_field, '')} {line}".strip()

    if current:
        papers.append(current)

    papers = papers[:limit]
    for paper in papers:
        paper["title"] = paper.get("title", "Untitled").replace("AutoNAS", PROJECT_TITLE)
        for key in ["authors", "year", "source", "abstract", "advantages", "datasets", "disadvantages", "results", "relevance"]:
            paper.setdefault(key, "Not specified")
    return papers


def generate_project_report() -> None:
    report = load_training_report()
    papers = load_paper_lines(16)
    models = report.get("all_models", [])
    best = report.get("best_model", {})
    task_type = report.get("task_type", "classification")
    best_architecture = format_architecture(best.get("architecture", {}))
    fastest_model = min(models, key=lambda model: model.get("training_time", float("inf")), default=None)
    smallest_model = min(models, key=lambda model: model.get("total_params", float("inf")), default=None)
    strongest_test_model = max(models, key=lambda model: model.get("final_accuracy", float("-inf")), default=None) if task_type == "classification" else None
    average_training_time = (
        sum(model.get("training_time", 0.0) for model in models if isinstance(model.get("training_time"), (int, float))) / len(models)
        if models else 0.0
    )

    doc = new_document("Project Report", f"{PROJECT_TITLE} – {FORMAL_SUBTITLE}")

    add_heading(doc, "ACKNOWLEDGEMENT", 1)
    add_long_paragraphs(doc, [
        "This report has been prepared by studying the actual implementation files, generated artifacts, interface templates, and literature notes available in the workspace. The goal of the report is to present the project in a professional but simple academic form without introducing invented claims.",
        "Gratitude is extended to the faculty guidance, academic reviewers, and the open-source software ecosystem that supported the development of this project. The practical availability of Flask, TensorFlow/Keras, Pandas, NumPy, scikit-learn, and python-docx made it possible to build both the software and its supporting documentation.",
        "The project also benefited from prior research studies in healthcare prediction and neural architecture search. Those studies helped shape the literature survey and guided the design decisions documented in the following chapters."
    ])

    add_heading(doc, "ABSTRACT", 1)
    add_long_paragraphs(doc, [
        f"{PROJECT_TITLE} is a web-based machine learning application developed to automate neural architecture search for tabular medical datasets. The system integrates a Flask interface, a preprocessing pipeline, and a TensorFlow/Keras training engine to create a practical end-to-end workflow from dataset upload to prediction.",
        "The project supports CSV and Excel file input, automatic task detection, preprocessing of numeric and categorical data, bounded generation of candidate dense neural networks, real-time progress reporting using Server-Sent Events, and persistent storage of learned artifacts for later inference.",
        f"The current saved workspace run corresponds to a {task_type} problem with {len(models)} candidate models. The best saved candidate is Candidate {best.get('candidate', 'N/A')} with architecture {best.get('architecture', {})}, validation metric {metric_percent(best.get('val_metric')) if task_type == 'classification' else best.get('val_metric')}, test accuracy {metric_percent(best.get('final_accuracy'))}, and {best.get('total_params', 'N/A')} trainable parameters. The project demonstrates that a controlled NAS workflow can be implemented as a usable academic software system for medical-style tabular prediction."
    ])
    doc.add_page_break()

    add_heading(doc, "LIST OF FIGURES", 1)
    for figure in [
        "Figure 3.1  Overall architecture of the proposed system",
        "Figure 3.2  Preprocessing flow for mixed-type input",
        "Figure 3.3  Candidate training and validation workflow",
        "Figure 3.4  NAS strategy execution overview",
        "Figure 4.1  Candidate-wise performance comparison",
        "Figure 4.2  User interface and system outputs",
    ]:
        add_paragraph(doc, figure)
    doc.add_page_break()

    add_heading(doc, "ABBREVIATIONS", 1)
    for item in [
        "API   : Application Programming Interface",
        "CSV   : Comma-Separated Values",
        "DL    : Deep Learning",
        "ML    : Machine Learning",
        "NAS   : Neural Architecture Search",
        "SDD   : Software Design Document",
        "SRS   : Software Requirements Specification",
        "SSE   : Server-Sent Events",
        "UI    : User Interface",
    ]:
        add_paragraph(doc, item)
    doc.add_page_break()

    add_heading(doc, "LIST OF TABLES", 1)
    for table in [
        "Table 2.1  Comparative analysis of existing systems",
        "Table 3.1  Search-space configuration and parameter constraints",
        "Table 3.2  Data preprocessing operations and purposes",
        "Table 3.3  API endpoints and purposes",
        "Table 3.4  NAS strategies and characteristics",
        "Table 4.1  Current run snapshot from training_report.json",
        "Table 4.2  Candidate-wise result summary",
        "Table 5.1  Project achievements summary",
    ]:
        add_paragraph(doc, table)
    doc.add_page_break()

    add_toc(doc, [
        ("ACKNOWLEDGEMENT", 0, "I"),
        ("ABSTRACT", 0, "II"),
        ("LIST OF FIGURES", 0, "III"),
        ("ABBREVIATIONS", 0, "IV"),
        ("LIST OF TABLES", 0, "V"),
        ("CHAPTER 1: INTRODUCTION", 0, "1"),
        ("1.1 Motivation", 1, "2"),
        ("1.2 Problem Statement", 1, "3"),
        ("1.3 Objectives", 1, "4"),
        ("1.4 Scope of the Project", 1, "5"),
        ("1.5 Proposed System Overview", 1, "6"),
        ("1.6 Key Features of the System", 1, "7"),
        ("1.7 Advantages of the Proposed System", 1, "8"),
        ("CHAPTER 2: LITERATURE SURVEY", 0, "10"),
        ("2.1 Introduction", 1, "10"),
        ("2.2 Existing Systems", 1, "11"),
        ("2.3 Comparative Analysis of Existing Systems", 1, "26"),
        ("2.4 Research Gaps Identified", 1, "27"),
        ("2.5 Summary", 1, "28"),
        ("CHAPTER 3: METHODOLOGY", 0, "29"),
        ("3.1 Different Phases of the System", 1, "29"),
        ("3.1.1 Data Collection and Input Format", 2, "29"),
        ("3.1.2 Data Preprocessing", 2, "30"),
        ("3.1.3 Feature Schema Generation", 2, "32"),
        ("3.1.4 Neural Architecture Search (NAS)", 2, "33"),
        ("3.1.5 Model Development", 2, "35"),
        ("3.1.6 Model Training", 2, "36"),
        ("3.1.7 Model Evaluation", 2, "37"),
        ("3.1.8 Best Model Selection", 2, "38"),
        ("3.1.9 Model Storage and Saving", 2, "39"),
        ("3.1.10 Prediction Module", 2, "40"),
        ("3.2 Overall Architecture of the System", 1, "41"),
        ("3.3 Neural Architecture Search Engine", 1, "42"),
        ("3.4 Model Structure", 1, "45"),
        ("3.5 Summary", 1, "48"),
        ("CHAPTER 4: RESULTS AND DISCUSSIONS", 0, "49"),
        ("4.1 Experimental Setup", 1, "49"),
        ("4.2 Dataset Description", 1, "50"),
        ("4.3 Candidate Models Generated", 1, "51"),
        ("4.4 Training and Validation Performance", 1, "52"),
        ("4.5 Best Model Analysis", 1, "54"),
        ("4.6 User Interface and System Output", 1, "56"),
        ("4.7 Discussion of Results", 1, "58"),
        ("4.8 Limitations of the System", 1, "59"),
        ("CHAPTER 5: CONCLUSION AND FUTURE SCOPE", 0, "60"),
        ("5.1 Conclusion", 1, "60"),
        ("5.2 Achievements of the Project", 1, "61"),
        ("5.3 Limitations", 1, "62"),
        ("5.4 Future Scope", 1, "63"),
        ("REFERENCES", 0, "64"),
        ("LIST OF PUBLICATIONS", 0, "65"),
    ])

    doc.add_page_break()
    add_heading(doc, "CHAPTER 1: INTRODUCTION", 1)
    add_heading(doc, "1.1 Motivation", 2)
    add_long_paragraphs(doc, [
        "Healthcare prediction tasks frequently depend on tabular datasets that contain a mix of demographic information, diagnostic indicators, laboratory measurements, and outcome labels. Such datasets are well suited for supervised learning, but selecting the best neural architecture for them is not straightforward.",
        "This project is motivated by the need to reduce manual trial-and-error in model design. Instead of asking a user to manually choose layer depth, hidden units, and activation functions repeatedly, the system explores these configurations automatically within a safe bounded search space.",
        "A second motivation is accessibility. By offering a browser-based workflow, the project helps students and evaluators observe how preprocessing, architecture search, training, reporting, and prediction are connected in a full software system.",
        "A third motivation is reproducibility. In academic machine learning work, promising results are often difficult to reproduce because preprocessing, label handling, train-validation-test splitting, and model configuration are scattered across notebooks or temporary scripts. The present project keeps those stages together in one controlled workflow and persists the resulting artifacts in the uploads directory for future verification.",
        "The project is also motivated by an educational need. Students who learn AutoML or neural architecture search often encounter theory-heavy explanations that do not show how search decisions interact with file uploads, backend concurrency, schema generation, and post-training prediction. This system narrows that gap by translating the concept of NAS into a complete working software product.",
        "Finally, the project responds to practical compute limitations. Rather than adopting extremely expensive NAS methods intended for image benchmarks and multi-GPU clusters, the implementation deliberately uses a constrained search space, early stopping, and low-depth feed-forward models so the search remains realistic in a lightweight academic environment."
    ])
    add_heading(doc, "1.2 Problem Statement", 2)
    add_long_paragraphs(doc, [
        "Neural architecture design for tabular medical prediction is often fragmented across multiple scripts and manual decisions. In many small projects, preprocessing logic, model training, evaluation, and prediction are not kept together in a reproducible manner.",
        "When the same preprocessing steps are not preserved exactly between training and inference, predictions on new inputs can become unreliable. Similarly, when candidate architectures are tested manually, comparison becomes difficult and time-consuming.",
        "The problem addressed by this project is the absence of a simple but integrated software platform that accepts real tabular data, preprocesses it consistently, performs bounded neural architecture search, saves artifacts, and supports prediction through a user interface.",
        "A related problem is the disconnect between model selection and deployment readiness in student projects. Even when a strong model is discovered, the absence of a saved preprocessor, target encoding logic, and input schema means the model cannot be applied safely to new records. The proposed system treats deployment-oriented concerns such as artifact persistence and input validation as part of the core design instead of an afterthought.",
        "The software must therefore solve both a machine learning problem and a systems problem. On the machine learning side it must prepare mixed-type tabular data, rank multiple neural candidates, and select the best model objectively. On the systems side it must expose routes, manage background execution, stream progress to the client, and provide downloadable outputs without blocking the web application."
    ])
    add_heading(doc, "1.3 Objectives", 2)
    add_numbered(doc, [
        "To build a web-based NAS workflow for tabular medical datasets.",
        "To support CSV and Excel dataset formats.",
        "To detect classification and regression tasks automatically.",
        "To implement random, evolutionary, and progressive search strategies.",
        "To provide real-time monitoring of training progress.",
        "To save reusable artifacts for download and later prediction.",
    ])
    add_report_table(doc, ["Objective", "Implemented Evidence in Workspace"], [
        ["Web-based ML workflow", "Flask routes in app.py render training and prediction pages"],
        ["Tabular data ingestion", "load_and_prepare_dataset() reads CSV, XLS, and XLSX"],
        ["Automated model search", "nas_engine.py implements random, evolutionary, and progressive search"],
        ["Live monitoring", "SSE queue and /stream route publish training events"],
        ["Artifact persistence", "best_pipeline.joblib, candidate_*.keras, and training_report.json are saved"],
    ])
    add_heading(doc, "1.4 Scope of the Project", 2)
    add_long_paragraphs(doc, [
        "The scope of the project includes dataset upload, preprocessing, task detection, candidate generation, model training, evaluation, best-model selection, artifact storage, and prediction support. These stages are exposed through a browser interface and a Flask backend.",
        "The project is intentionally restricted to dense neural networks for structured tabular data. It does not attempt to address image NAS, text NAS, distributed cluster execution, or hospital-grade deployment controls.",
        "This scope makes the project suitable for academic evaluation and reproducible prototype experimentation while still covering a complete end-to-end machine learning workflow.",
        "Within that scope, the software addresses both training-time and inference-time concerns. The training portion covers data acceptance, preprocessing, architecture generation, controlled optimization, and reporting. The inference portion covers schema discovery, dynamic form generation, user input coercion, model loading, and prediction response formatting.",
        "The project scope also includes engineering safeguards. Examples visible in the codebase include range clamping for user-supplied hyperparameters, a one-million-parameter cap in the NAS engine, and early stopping to limit unnecessary training. These measures show that the implementation values stability and responsiveness in addition to predictive performance."
    ])
    add_heading(doc, "1.5 Proposed System Overview", 2)
    add_long_paragraphs(doc, [
        "The proposed system consists of a frontend interface, a Flask-based application layer, a preprocessing pipeline, and a neural architecture search engine. The user uploads a dataset and chooses a NAS strategy from the training page. The backend stores the file, starts a background worker, prepares the data, and executes the selected search process.",
        "During training, progress updates are streamed to the interface through Server-Sent Events. After training, the system stores the best model, a preprocessing bundle, and a JSON report. Those saved artifacts are then reused by the prediction module.",
        "As a result, the software behaves as a complete training-and-inference platform rather than a one-time experiment script.",
        "The overall process can be understood as a closed loop. Input data enters the system through a controlled upload route, is standardized through a preprocessing transformer, is used to evaluate candidate architectures, and finally produces both quantitative metrics and reusable files. Those same files later feed the prediction route, ensuring that the user-facing prediction workflow is directly tied to the training-time configuration.",
        "The reportable outcome of the currently saved run reflects this integrated design. The workspace stores five candidate models, a best pipeline bundle, and a report whose best candidate is Candidate 3 with a validation metric of 88.89%, a test accuracy of 79.63%, and 2,658 parameters. These saved values make the project more than a theoretical prototype because they provide verifiable outputs for academic assessment."
    ])
    add_heading(doc, "1.6 Key Features of the System", 2)
    add_bullets(doc, [
        "CSV, XLS, and XLSX upload support",
        "Automatic task detection",
        "Random, evolutionary, and progressive NAS",
        "Server-Sent Event based progress monitoring",
        "Artifact download support",
        "Schema-driven prediction interface",
    ])
    add_long_paragraphs(doc, [
        "Each feature contributes to a unified learning experience. File-format flexibility lowers entry barriers for users who may have exported data from spreadsheets or medical records systems. Automatic task detection reduces configuration errors when labels are stored numerically. The combination of multiple search strategies allows the same application to demonstrate different NAS philosophies without changing the surrounding workflow.",
        "The progress-streaming feature is especially important in a web-based academic tool because model training is no longer a black box. Users can see that work is proceeding, observe candidate-level updates, and connect system design choices with runtime behavior. The dynamic prediction schema likewise improves usability by generating input fields from the saved training metadata instead of relying on hard-coded forms.",
        "Another practical feature is the system's artifact-centered workflow. The user is not left with a temporary result displayed on the screen; instead, the software persists the selected model, the candidate models, the preprocessing bundle, and the machine-readable training report. This approach supports later verification, demonstration, and reuse, which are all important in an academic project review setting.",
        "The feature set also reveals a deliberate balance between flexibility and safety. The user can influence the search mode and related settings, but the backend still clamps unsafe values and protects the system from unrealistic model sizes. In this way, the software is not simply permissive; it is intentionally guided toward configurations that fit the project's educational and operational constraints."
    ])
    add_heading(doc, "1.7 Advantages of the Proposed System", 2)
    add_long_paragraphs(doc, [
        "The system reduces manual architecture selection effort by automating candidate exploration under controlled limits.",
        "It improves reproducibility because the preprocessing bundle and training report are saved together with the trained model.",
        "It improves usability because users can observe the workflow through a browser without needing to understand every internal training detail in advance.",
        "It also improves maintainability. The code separates responsibilities across `app.py`, `data_pipeline.py`, and `nas_engine.py`, which makes the implementation easier to document, test conceptually, and extend in future semesters.",
        "Another advantage is traceability. Because the training output is summarized in a structured JSON file and linked to downloadable artifacts, the basis for selecting the best model remains transparent. This is a stronger academic position than reporting a final accuracy value without preserving how that value was produced.",
        "The final advantage is practicality. Even though advanced NAS literature often assumes expensive hardware, the present implementation shows that a bounded search strategy can still deliver educational and functional value on a small system using CPU-friendly settings.",
        "There is also an advantage in pedagogical clarity. Because the system keeps its major operations visible through routes, saved files, and predictable module boundaries, it becomes easier for a reviewer to understand how a dataset moves from upload to preprocessing, from preprocessing to model search, and from search to prediction. Many academic projects produce results, but fewer make the complete reasoning chain easy to inspect.",
        "A further advantage lies in the system's adaptability. Although the motivating use case is medical tabular prediction, the design does not hard-code a single disease label or a fixed column set. The same workflow can be reused for similar structured datasets with minimal adjustment, making the project a reusable academic platform rather than a one-dataset script."
    ])

    doc.add_page_break()
    add_heading(doc, "CHAPTER 2: LITERATURE SURVEY", 1)
    add_heading(doc, "2.1 Introduction", 2)
    add_long_paragraphs(doc, [
        "The literature survey for this project spans healthcare prediction studies and neural architecture search research. Healthcare studies provide domain relevance and demonstrate how structured medical datasets are used in predictive modeling. NAS studies provide the methodological basis for automated architecture exploration.",
        "Together, these two literature streams justify the design of a bounded, tabular-data-focused NAS system intended for healthcare-style prediction workflows."
    ])
    add_heading(doc, "2.2 Existing Systems", 2)
    selected_papers = papers[:16]
    add_paragraph(doc, "This section uses a compact review of all sixteen summarized studies collected for the project so that the chapter covers the complete literature set while remaining focused on the most relevant ideas for tabular healthcare prediction and neural architecture search.")
    for index, paper in enumerate(selected_papers, start=1):
        title = paper.get('title', f'Paper {index}')
        title_lower = title.lower()
        focus = "Healthcare prediction" if "heart" in title_lower or "disease" in title_lower or "medical" in title_lower else "NAS / AutoML"
        add_heading(doc, f"2.2.{index} {title}", 3)
        add_paragraph(
            doc,
            f"This {focus.lower()} study by {paper.get('authors', 'the listed authors')} ({paper.get('year', 'n.d.')}) reports {paper.get('results', 'key results not specified')}. Its main contribution is {paper.get('advantages', 'not specified')}. For the current project, the paper is relevant because {paper.get('relevance', 'it informs the design direction of the proposed system')}. The main limitation noted in the summary is {paper.get('disadvantages', 'not specified')}, which reinforces the need for a bounded, reproducible, tabular-focused NAS workflow."
        )
    add_heading(doc, "2.3 Comparative Analysis of Existing Systems", 2)
    comparative_rows = []
    for paper in selected_papers:
        comparative_rows.append([
            paper.get("title", "Paper")[:32],
            paper.get("year", "NA"),
            "Healthcare" if "heart" in paper.get("title", "").lower() else "NAS/AutoML",
            paper.get("results", "Not specified")[:36],
            "High",
        ])
    add_report_table(doc, ["Paper", "Year", "Category", "Key Result", "Relevance"], comparative_rows)
    add_long_paragraphs(doc, [
        "The compact comparison shows a clear divide between domain-oriented healthcare prediction studies and methodology-oriented NAS or AutoML studies. Healthcare papers provide problem relevance and common benchmark framing, whereas NAS papers explain how search strategy, search space, and computational cost affect automated architecture selection.",
        "The proposed project sits between these two streams. It adopts the practical medical motivation of structured disease prediction while borrowing the search-oriented thinking of NAS research. This hybrid position explains why the project focuses on a compact search space, consistent preprocessing, and artifact-preserving execution rather than on very large benchmark-driven experimentation."
    ])
    add_heading(doc, "2.4 Research Gaps Identified", 2)
    add_long_paragraphs(doc, [
        "Many healthcare prediction studies use fixed model families chosen manually and do not automate architecture exploration.",
        "Many NAS studies focus on image benchmarks and assume higher computational resources than a lightweight tabular medical workflow requires.",
        "The literature also shows a gap between algorithm-focused research and reusable software systems. Even strong studies do not always preserve preprocessing, trained artifacts, and inference metadata in a way that supports later demonstration and prediction reuse.",
        "The project addresses these gaps by integrating tabular preprocessing, bounded architecture search, result tracking, and prediction support in one practical application."
    ])
    add_heading(doc, "2.5 Summary", 2)
    add_long_paragraphs(doc, [
        "The literature supports the need for a project that combines the practical relevance of healthcare prediction with the automation benefits of NAS. The proposed system is designed as that bridge.",
        "By limiting the review to the most representative studies, the chapter remains focused on the ideas that matter most for this project: tabular medical relevance, compact model design, resource-aware search, and the value of preserving artifacts for later use."
    ])

    doc.add_page_break()
    add_heading(doc, "CHAPTER 3: METHODOLOGY", 1)
    add_heading(doc, "3.1 Different Phases of the System", 2)
    add_heading(doc, "3.1.1 Data Collection and Input Format", 3)
    add_long_paragraphs(doc, [
        "The workflow begins when a user uploads a CSV, XLS, or XLSX file through the training interface. The file is sanitized and stored with a timestamp prefix inside the uploads directory.",
        "The last column is treated as the target variable while the remaining columns are treated as features. This allows the system to remain generic across tabular datasets.",
        "This design choice reduces the amount of manual configuration required from the user. Instead of asking for separate schema files or column mappings, the system uses a simple and transparent convention: all columns except the last are predictors and the last column is the prediction target. For academic datasets arranged in spreadsheet form, this convention is practical and easy to explain.",
        "The upload handling in the Flask layer also serves an engineering role. Filenames are sanitized before storage, and the application creates timestamped names to reduce collisions between runs. These operational details matter because a web-based training system must manage files safely and predictably, not merely fit a model.",
        "The accepted file formats are intentionally chosen to match what students and evaluators are likely to possess. CSV files are common for exported datasets, while XLS and XLSX support spreadsheet-based preparation. By accommodating both, the software reduces the friction between raw data collection and machine learning experimentation.",
        "This phase also establishes the accountability of the workflow. Once the file is stored, subsequent preprocessing and training operations are attached to an identifiable uploaded artifact rather than to a hidden in-memory table alone. That linkage is beneficial during evaluation because it improves traceability across the full experimental lifecycle."
    ])
    add_bullets(doc, ["CSV File Input", "Excel File Input"])
    add_heading(doc, "3.1.2 Data Preprocessing", 3)
    add_long_paragraphs(doc, [
        "The preprocessing stage separates numeric and categorical columns. Numeric columns are imputed using median values and scaled using StandardScaler. Categorical columns are imputed using mode values and one-hot encoded.",
        "The dataset is split into training, validation, and test sets using an effective 60/20/20 distribution. This supports validation-driven model selection and separate test evaluation.",
        "The choice of median imputation for numeric variables is robust against skewed measurements and outliers that often occur in clinical or semi-clinical tabular records. Mode imputation for categorical variables preserves the discrete nature of symbols such as chest-pain categories, yes/no diagnostic indicators, or sex-coded variables.",
        "Standardization of numeric variables helps stabilize optimization in dense neural networks because features with widely different scales can otherwise dominate the gradient signal. One-hot encoding likewise converts categorical values into a representation that a feed-forward network can process without imposing false ordinal meaning on symbolic labels.",
        "The preprocessing pipeline is implemented with scikit-learn `Pipeline` and `ColumnTransformer` objects. This is significant because it packages all transformation logic into a single serializable object, which is later saved with joblib and reused during prediction. In other words, methodology here is not just conceptual; it is embodied in a reusable software artifact."
    ])
    add_bullets(doc, [
        "Missing Value Handling",
        "Numeric Feature Scaling",
        "Categorical Encoding (One-Hot Encoding)",
        "Train / Validation / Test Split (60/20/20)",
    ])
    add_report_table(doc, ["Operation", "Technique", "Purpose"], [
        ["Missing numeric values", "Median imputation", "Preserve rows and stabilize numeric input"],
        ["Missing categorical values", "Mode imputation", "Retain categorical structure"],
        ["Numeric scaling", "StandardScaler", "Normalize feature ranges"],
        ["Categorical conversion", "OneHotEncoder", "Convert categories into model-ready columns"],
    ])
    add_long_paragraphs(doc, [
        "The sequence of preprocessing operations is not arbitrary. Imputation occurs before scaling and encoding because the downstream transformations require complete input values. The structured ordering of these steps within a scikit-learn pipeline prevents accidental inconsistency between training and prediction time, which is one of the common failure points in small machine learning projects.",
        "The preprocessing design also reflects a principle of practical robustness. Medical-style datasets often contain missing values, mixed measurement scales, and categorical indicators. By explicitly separating numeric and categorical pathways, the system avoids the error of applying a single generic transformation to fundamentally different data types.",
        "Another benefit of the chosen preprocessing pipeline is reusability. Once fitted, the transformer becomes a portable representation of the exact feature engineering used in training. This makes the project stronger than a script that computes temporary transformations ad hoc and then loses the mapping needed for later inference."
    ])
    add_heading(doc, "3.1.3 Feature Schema Generation", 3)
    add_long_paragraphs(doc, [
        "The preprocessing module builds a feature schema for all input columns. Numeric features record summary information and categorical features record allowed choices.",
        "This schema is later used to generate the prediction form dynamically, ensuring that new inputs follow the same structural assumptions used during training.",
        "For numeric features, the schema includes min, max, mean, and an example value. For categorical features, it includes the allowable choices and an example. These values make the prediction page more informative and reduce the chance of malformed input.",
        "From a software-engineering perspective, feature schema generation acts as the contract between the training subsystem and the inference interface. It is one of the strongest examples in the project of how machine learning metadata can be transformed into usable frontend behaviour."
    ])
    add_bullets(doc, ["Numeric Feature Metadata", "Categorical Feature Choices", "Dynamic Form Generation"])
    add_long_paragraphs(doc, [
        "Feature schema generation can also be understood as a documentation mechanism embedded inside the software. It captures practical information about each input column in a form that the frontend can use directly. This transforms descriptive metadata into an executable interface contract.",
        "Because the prediction form is created from saved schema information, the software reduces the risk of interface drift. A manually coded form can become outdated when the training dataset changes, but a schema-driven form remains aligned with the actual training inputs. This is especially valuable in an academic setting where multiple datasets may be tried during demonstration."
    ])
    add_heading(doc, "3.1.4 Neural Architecture Search (NAS)", 3)
    add_long_paragraphs(doc, [
        "The NAS engine explores dense feed-forward architectures in a bounded search space. Candidate models vary in hidden-layer depth, unit counts, and activation function while remaining suitable for tabular data.",
        "Three strategies are implemented: random search, evolutionary search, and progressive search. Each strategy operates within practical runtime and parameter constraints.",
        "The search space is intentionally compact: from one to four hidden layers, with units drawn from {16, 32, 64, 128} and activations drawn from {relu, tanh}. This limited design is not a weakness; it is a deliberate response to the system context. A small but meaningful search space is more appropriate for tabular educational experimentation than very large controller-based or differentiable spaces designed for benchmark-heavy settings.",
        "The NAS engine also enforces a one-million-parameter cap. Candidates that exceed this size are skipped, protecting the system from architectures that would add computational cost without fitting the lightweight web-based objective. In this sense, the methodology explicitly blends search freedom with runtime governance.",
        "Search strategies are implemented as alternative ways of traversing the same bounded space. Random search provides a baseline of independent sampling, evolutionary search introduces mutation-based refinement, and progressive search expands model depth only when it seems worthwhile. This allows the project to compare algorithmic styles while keeping the rest of the workflow fixed."
    ])
    add_bullets(doc, [
        "Search Space Definition",
        "Number of Layers (1–4)",
        "Units per Layer (16, 32, 64, 128)",
        "Activation Functions (ReLU, Tanh)",
        "Random Search",
        "Evolutionary Search",
        "Progressive Search",
    ])
    add_report_table(doc, ["Search Component", "Implemented Choice"], [
        ["Depth", "1 to 4 hidden layers"],
        ["Units", "16, 32, 64, 128"],
        ["Activations", "relu, tanh"],
        ["Parameter guard", "Skip candidates above 1,000,000 parameters"],
    ])
    add_long_paragraphs(doc, [
        "The bounded search space should be interpreted as a strategic simplification rather than a reduction in seriousness. A useful academic prototype does not need to explore every possible topology. Instead, it needs to search a meaningful subset of architectures in a way that produces understandable results and a reliable workflow.",
        "The use of three search strategies within the same implementation creates an additional layer of educational value. The project does not merely say that NAS can be automated; it demonstrates that automation itself can follow different philosophies, ranging from purely random exploration to mutation-guided refinement and staged architectural growth.",
        "The parameter cap reinforces the project's concern for responsible experimentation. It protects the system from models whose computational cost would be disproportionate to the expected benefit in a small structured-data setting. In this sense, the search engine encodes methodological discipline directly into its runtime behaviour."
    ])
    add_heading(doc, "3.1.5 Model Development", 3)
    add_long_paragraphs(doc, [
        "Candidate models are built using the Keras Functional API so that architectures can be generated dynamically while preserving a consistent interface.",
        "The output layer is task-aware: softmax is used for classification and a single linear unit is used for regression.",
        "The Functional API is important here because the number and size of hidden layers are not fixed in advance. The model-building function receives an architecture dictionary and constructs the dense stack programmatically. This makes the codebase faithful to the concept of search: each candidate architecture is a first-class configuration, not a manual rewrite of a hard-coded network.",
        "Using a task-aware output layer also reinforces the software's generic design. The same training system can adapt to classification or regression without changing the surrounding upload or preprocessing workflow, because the prepared data object and model builder carry the task information forward."
    ])
    add_bullets(doc, ["Dense Neural Network Construction", "Activation Function Assignment", "Output Layer Design", "Softmax (Classification)", "Linear (Regression)"])
    add_long_paragraphs(doc, [
        "Model development in this project is tightly coupled to the architecture representation stored in Python dictionaries. This design keeps the search logic and the model-building logic aligned: the same structure that is sampled, mutated, or grown by the NAS engine is the structure that determines the Keras graph. Such alignment reduces the possibility of mismatch between search metadata and actual instantiated networks.",
        "The choice of dense neural networks also reflects a realistic view of the underlying data. After preprocessing, the features exist as a flattened numerical vector, which makes feed-forward layers a natural and computationally efficient choice. The methodology therefore adapts the model family to the data representation rather than forcing a more complex architecture for the sake of novelty alone."
    ])
    add_heading(doc, "3.1.6 Model Training", 3)
    add_long_paragraphs(doc, [
        "Model training uses TensorFlow/Keras with tf.data based batching and prefetching. Each candidate uses Adam optimizer with a learning rate of 0.001.",
        "An early stopping callback monitors validation loss and reduces unnecessary computation.",
        "Training is intentionally conservative. The code clamps epochs to a small range, uses a modest learning rate, and clears TensorFlow sessions between candidates. These decisions reduce memory pressure and keep the application responsive enough for a web-hosted educational demonstration.",
        "The training subsystem is not isolated from the user interface. Through callback-driven event publishing, intermediate information is pushed onto the SSE queue so that the browser can visualize progress. This creates a methodology in which optimization and transparency are linked rather than treated as separate concerns."
    ])
    add_bullets(doc, ["TensorFlow / Keras Framework", "Adam Optimizer (Learning Rate = 0.001)", "Early Stopping Mechanism"])
    add_long_paragraphs(doc, [
        "The training configuration is intentionally modest but pedagogically strong. It is sufficient to show how a complete model-selection loop works, while remaining fast enough to support live observation through the web interface. In this way, the system privileges transparency and repeatability over brute-force optimization.",
        "Using a common optimizer such as Adam also improves interpretability of the experiment. Reviewers familiar with neural networks can understand the training setup without needing to decode an exotic optimization procedure. This helps keep the focus of the project on the NAS workflow and the integration of the software components.",
        "The combination of tf.data batching, early stopping, and per-candidate cleanup demonstrates that the implementation takes runtime management seriously. These measures help ensure that several candidate models can be trained sequentially without gradually exhausting resources or making the web application unresponsive."
    ])
    add_heading(doc, "3.1.7 Model Evaluation", 3)
    add_long_paragraphs(doc, [
        "Candidates are evaluated on validation and test sets. Accuracy is used for classification tasks and mean squared error is used for regression tasks. Loss values are also preserved.",
        "Keeping validation and test metrics separate helps maintain correct model-selection practice.",
        "This distinction is academically important. Validation performance determines which architecture is preferred during search, while the test set provides a more honest final estimate after that choice has been made. The current saved run illustrates this point clearly because the model with the best validation performance is not simply the one with the visually highest test accuracy; instead, it is the candidate judged strongest under the project's ranking rule.",
        "In classification mode, the report stores both the validation metric and final accuracy, along with final loss. This allows discussion of trade-offs such as compactness versus generalization, or selection metric versus final observed test outcome."
    ])
    add_bullets(doc, ["Training Loss", "Validation Loss", "Validation Accuracy", "Test Accuracy"])
    add_long_paragraphs(doc, [
        "Evaluation in the project is not limited to declaring a winner. It provides the evidential basis for comparing architectures, discussing trade-offs, and justifying why a particular candidate should be reused for prediction. This broader role is important because NAS is only meaningful when the comparison mechanism is methodologically sound.",
        "The stored metrics also enable retrospective analysis. Once the training report is saved, the project team can revisit how candidate complexity, training time, validation behaviour, and final loss related to each other. This transforms the report from a simple output artifact into a source of empirical interpretation."
    ])
    add_heading(doc, "3.1.8 Best Model Selection", 3)
    add_long_paragraphs(doc, [
        "The best model is selected based on validation performance, not final test performance. This avoids test leakage during model ranking.",
        "For classification, higher validation accuracy is preferred. For regression, lower validation MSE is treated as a better ranking signal.",
        "This choice reinforces methodological discipline. If the test set were used for selection, the same data would influence both model choice and final performance reporting, weakening the credibility of the result. The current software avoids that mistake by computing a ranking signal called `val_performance` and selecting the best candidate using that value.",
        "The selected result in the saved workspace is Candidate 3, which achieved the strongest validation metric among the five candidates while remaining compact in parameter count. That combination makes it a credible best-model choice under the project's stated design principles."
    ])
    add_bullets(doc, ["Selection using Validation Accuracy", "Selection using Validation MSE"])
    add_long_paragraphs(doc, [
        "Best-model selection is one of the most sensitive phases in an AutoML-style workflow because an error here can invalidate the credibility of the entire experimental process. The project addresses this by making the ranking criterion explicit and by preserving the winning candidate's metadata inside the training report.",
        "The preserved selection trail also benefits future documentation. A reviewer can inspect which candidate won, what architecture it used, how it performed on validation and test sets, and why that outcome was reasonable. This traceability is especially valuable when several candidates have broadly similar final accuracy but differ in compactness or validation stability."
    ])
    add_heading(doc, "3.1.9 Model Storage and Saving", 3)
    add_long_paragraphs(doc, [
        "Each candidate model is stored as a .keras file. The fitted preprocessor, label encoder, feature schema, and related metadata are stored as a joblib bundle. The run report is stored as JSON.",
        "This persistence design allows the system to reuse exactly the same preprocessing and model artifacts during later prediction.",
        "Artifact persistence is one of the strongest practical aspects of the project. Instead of saving only the best neural weights, the software preserves the context required to interpret and reuse those weights. This includes feature ordering, categorical expansion behaviour, class-label mapping, and summary metrics.",
        "The saved files in the uploads directory therefore function as evidence of both process and outcome. They also make the report defensible because reported results can be traced to specific generated artifacts rather than to undocumented transient memory objects."
    ])
    add_bullets(doc, ["Saving Model (.keras)", "Saving Preprocessor (.joblib)", "Training Report Storage"])
    add_long_paragraphs(doc, [
        "The storage strategy effectively turns the project into a reusable experiment package. A future user does not need to repeat training merely to demonstrate inference, because the required model and preprocessing context are already available in the saved outputs. This separation between training time and demonstration time makes the system more practical for academic evaluation schedules.",
        "Persistent storage also strengthens reproducibility claims. When model artifacts are saved explicitly, report statements about the best model are anchored in files that can be inspected, downloaded, and reused. This is more convincing than describing a best model whose exact weights and preprocessing history are no longer available."
    ])
    add_heading(doc, "3.1.10 Prediction Module", 3)
    add_long_paragraphs(doc, [
        "The prediction module retrieves the saved schema and best model to accept new user input. Numeric and categorical inputs are validated, transformed with the stored preprocessor, and passed to the saved model.",
        "Classification predictions return label probabilities and regression predictions return numeric values.",
        "In the Flask application, the best model is cached in memory to avoid loading the `.keras` file repeatedly on every prediction request. This reduces repeated overhead and demonstrates how model-serving considerations can be integrated even in a small academic system.",
        "The prediction module is significant because it closes the loop between search and deployment. A searched model is not useful if it cannot later accept new records through a stable interface. By combining schema-aware input handling, preprocessing reuse, and model caching, the project turns search output into a usable service."
    ])
    add_bullets(doc, ["Prediction Schema API", "User Input Handling", "Data Preprocessing for Prediction", "Output Prediction (Class / Probability / Value)"])
    add_long_paragraphs(doc, [
        "The prediction workflow is important because it converts the outcome of NAS into an actionable software capability. Without an inference module, the system would stop at architecture comparison and would not demonstrate practical usability. By adding a prediction route and schema-driven form, the software shows how model search results can become part of a functioning application.",
        "The caching of the best model is a small but meaningful serving optimization. It reflects awareness that deployment-oriented systems should not repeatedly incur unnecessary model-loading overhead. Even though this project is academic in scale, it still incorporates that operational mindset into its design."
    ])
    add_heading(doc, "3.2 Overall Architecture of the System", 2)
    add_long_paragraphs(doc, [
        "The system follows a layered architecture consisting of presentation, application, and machine learning layers. The presentation layer contains the training and prediction pages. The application layer manages routes, state, downloads, and streaming. The ML layer performs preparation, search, training, and evaluation.",
        "This separation improves clarity and maintainability.",
        "The presentation layer is implemented through HTML templates that expose two clear workflows: training and prediction. The application layer coordinates requests and background processing using Flask routes, queue-driven SSE, and a runtime state dataclass protected by locks. The machine learning layer contains the transformation and optimization logic that is intentionally separated into dedicated modules.",
        "This layered pattern helps the report explain the system at different levels. Non-technical stakeholders can understand the page-level flow, while technical readers can follow how route handlers call the preprocessing module and NAS engine. Such architectural separation is especially valuable when the project is assessed simultaneously for software design and machine learning capability."
    ])
    add_report_table(doc, ["Layer", "Main Elements", "Responsibilities"], [
        ["Presentation", "templates/index.html, templates/predict.html", "Collect user input, display progress, render results"],
        ["Application", "Flask routes, RuntimeState, SSE queue", "Handle requests, manage background execution, serve artifacts"],
        ["ML Processing", "data_pipeline.py", "Prepare data, detect task, build schema, split datasets"],
        ["NAS Engine", "nas_engine.py", "Generate architectures, train candidates, rank and summarize results"],
        ["Persistence", "uploads directory", "Store models, pipeline bundle, and JSON reports"],
    ])
    add_report_table(doc, ["Endpoint", "Method", "Purpose"], [
        ["/", "GET", "Open the training page"],
        ["/predict_page", "GET", "Open the prediction page"],
        ["/upload", "POST", "Receive dataset and NAS settings, then start training"],
        ["/stream", "GET", "Stream live training events via SSE"],
        ["/prediction_schema", "GET", "Expose saved feature metadata for dynamic form generation"],
        ["/predict", "POST", "Run inference using saved artifacts"],
        ["/download_best", "GET", "Download the best model file"],
        ["/download_report", "GET", "Download the training report JSON"],
        ["/download_candidate/<id>", "GET", "Download a specific candidate model"],
    ])
    add_report_table(doc, ["Architectural Concern", "Current Design Response", "Academic Significance"], [
        ["Responsiveness during training", "Background worker plus SSE stream", "Demonstrates that ML workflows can coexist with interactive web behavior"],
        ["Consistency between training and prediction", "Saved pipeline bundle and schema reuse", "Supports reproducibility and dependable inference"],
        ["Traceable model selection", "Structured training_report.json and candidate artifacts", "Provides evidence for academic evaluation and auditing"],
        ["Controlled runtime complexity", "Clamped settings and parameter guard", "Shows that engineering safeguards are part of the design"],
        ["Extensibility", "Separated modules for app, pipeline, and NAS engine", "Allows future enhancement without rewriting the whole system"],
    ])
    add_long_paragraphs(doc, [
        "Architecturally, the project succeeds because each layer has a clear responsibility while still participating in a common workflow. The presentation layer gathers intent, the application layer manages execution and state, the machine learning layer produces candidate outcomes, and the persistence layer preserves those outcomes for later use. This creates a coherent system narrative that is easy to explain in academic review.",
        "The architecture also supports incremental enhancement. Because data preparation, model search, and web serving are already separated into modules, future improvements can be introduced with less risk of destabilizing the entire application. This modularity is not incidental; it is one of the software design strengths of the project.",
        "The runtime architecture adds another layer of maturity. Training is delegated to a background worker rather than being executed directly in the request-response path, and progress information is forwarded to the interface through a queue-backed SSE mechanism. This keeps the application responsive and demonstrates that the project addresses operational behavior as well as algorithmic behavior.",
        "The architecture is also strengthened by the use of structured shared state. Runtime metadata such as candidate outcomes, artifact paths, feature schema, and task type are grouped into a single state container protected by locks. This organization reduces ambiguity, makes the code easier to reason about, and improves the reportability of the system design."
    ])
    add_heading(doc, "3.3 Neural Architecture Search Engine", 2)
    add_heading(doc, "3.3.1 Random Search Strategy", 3)
    add_long_paragraphs(doc, [
        "Random search samples independent candidates from the search space. It is simple, direct, and useful as a baseline strategy.",
        "Its methodological value lies in its neutrality. Because each architecture is sampled independently, random search provides a clear baseline against which more structured strategies can be discussed. In an academic setting, this is useful because it helps students understand whether later improvements are due to the search strategy or merely to the existence of multiple trials.",
        "In the present project, random search is especially suitable when quick exploration is preferred over guided iteration. It also aligns well with the bounded search space, where even simple independent sampling can discover practical compact models."
    ])
    add_heading(doc, "3.3.2 Evolutionary Search Strategy", 3)
    add_long_paragraphs(doc, [
        "Evolutionary search keeps stronger candidates and produces new ones through mutation, allowing guided exploration.",
        "The evolutionary idea is implemented in a lightweight way appropriate to the project scale. Rather than building a full research-grade evolutionary framework, the code retains better-performing candidates and mutates their unit configurations or activation choice. This makes the strategy understandable and computationally manageable.",
        "The educational benefit of this strategy is that it introduces the principle of search guidance: good solutions can influence the next generation. This contrasts with random search and illustrates why different NAS strategies exist in the literature."
    ])
    add_heading(doc, "3.3.3 Progressive Search Strategy", 3)
    add_long_paragraphs(doc, [
        "Progressive search grows model depth only when validation performance improves, helping conserve compute resources.",
        "This strategy is highly compatible with tabular data experiments because deeper networks are not always better for structured datasets. By beginning with simpler architectures and growing only when justified, the method operationalizes the idea of complexity control.",
        "Progressive search also reinforces one of the project's central design themes: bounded improvement. The system does not assume that more layers, more parameters, or more training automatically imply a better result. Instead, it treats complexity as something that must earn its place through validation gains."
    ])
    add_report_table(doc, ["Strategy", "Strength", "Limitation"], [
        ["Random Search", "Simple and diverse", "No guided reuse of prior candidates"],
        ["Evolutionary Search", "Guided improvement", "More coordination overhead"],
        ["Progressive Search", "Compute-aware growth", "Can stop early on local plateaus"],
    ])
    add_long_paragraphs(doc, [
        "Together, these strategies make the NAS engine more than a single heuristic. They turn it into a comparative educational framework within the same project. A student can observe that different search styles explore the same bounded model family in different ways and may favour different compromises between diversity, guidance, and computational restraint.",
        "This multi-strategy design also strengthens the report's methodological richness. Instead of saying that the project performs NAS in only one predetermined way, it can demonstrate that search itself is a design choice and that software architecture can expose those choices as configurable operational modes.",
        "The engine also shows how backend complexity can be hidden behind an approachable interface. The user chooses a search mode at a high level, but internally that choice determines how candidates are sampled, mutated, or expanded. This separation between user simplicity and backend richness is one of the reasons the project works well as an academic demonstration.",
        "Another important aspect is the preservation of architecture metadata. Search does not end when a model is trained; it produces a structured record of how that model was configured and how it performed. This makes the engine useful not only for finding a winner but also for supporting later explanation and comparison.",
        "Seen from a systems perspective, the NAS engine acts as the analytical core of the project. It converts prepared data into a set of comparable experimental outcomes and then hands those outcomes to the reporting and persistence layers. This role makes it central not only to model quality but also to the interpretability of the overall application."
    ])
    add_report_table(doc, ["Search Strategy", "Internal Behavior", "Practical Outcome in Project Context"], [
        ["Random", "Independent architecture sampling", "Useful baseline and quick diversity"],
        ["Evolutionary", "Mutation of stronger candidates", "Illustrates guided improvement under constraints"],
        ["Progressive", "Depth growth only when useful", "Supports complexity control for tabular tasks"],
    ])
    add_heading(doc, "3.4 Model Structure", 2)
    add_heading(doc, "3.4.1 Dense Neural Network Architecture", 3)
    add_long_paragraphs(doc, [
        "The model family used in this project is a dense feed-forward neural network suitable for structured tabular features.",
        "Dense layers are a practical choice for this data type because the preprocessed feature matrix is already a flat vector after scaling and one-hot encoding. Unlike image or sequence data, the tabular representation does not require convolutional or recurrent structure to become usable.",
        f"The best saved architecture in the current workspace can be summarized as {best_architecture}. This compact result supports the design assumption that tabular problems often benefit from moderate-capacity networks rather than excessively deep or wide designs."
    ])
    add_heading(doc, "3.4.2 Activation Functions Used", 3)
    add_long_paragraphs(doc, [
        "The search space currently uses relu and tanh activation functions for hidden layers.",
        "ReLU provides sparse activation and efficient optimization behaviour, while tanh offers a bounded nonlinearity that can be useful when the signal benefits from centering. Restricting the activation choices to two common functions keeps the search interpretable and reduces unnecessary combinatorial growth.",
        "The current best model uses tanh activations, which is a meaningful outcome because it shows the search process is not biased toward a single default. Instead, the selected activation emerges from comparative validation performance."
    ])
    add_heading(doc, "3.4.3 Output Layer Configuration", 3)
    add_long_paragraphs(doc, [
        "Classification uses softmax output while regression uses a single linear output neuron.",
        "This distinction keeps the model semantically aligned with the problem type detected earlier in the preprocessing stage. In classification, the output layer represents class probabilities, enabling the prediction route to return both label and confidence-style information. In regression, a scalar linear output supports continuous estimation.",
        "Because the output configuration is derived from the prepared data object rather than manually toggled in the UI, the design minimizes mismatch between dataset properties and network behaviour."
    ])
    add_long_paragraphs(doc, [
        "The model structure chapter as a whole demonstrates that simplicity and adaptability can coexist. The networks are not architecturally exotic, yet they are dynamic, task-aware, and selected through a structured search process. This makes them appropriate for the project's goal of serving as a full academic demonstration of NAS for tabular data.",
        "The evidence from the saved run supports this interpretation. A compact architecture with tanh activations emerged as the preferred candidate, indicating that the system's constrained search space is capable of producing nuanced outcomes rather than merely reproducing a fixed default configuration."
    ])
    add_heading(doc, "3.5 Summary", 2)
    add_long_paragraphs(doc, [
        "The methodology provides a consistent path from raw upload through preprocessing, model search, evaluation, artifact storage, and prediction reuse.",
        "Importantly, each methodological stage has a direct implementation counterpart in the workspace. That traceability strengthens the academic legitimacy of the report because the described process can be mapped to actual code, saved files, and generated outputs.",
        "As a result, Chapter 3 does not describe an abstract pipeline disconnected from implementation. It documents a method that is both technically coherent and operationally realized, which is one of the strongest features of the overall project.",
        "The chapter also clarifies why the project should be treated as a full system rather than as a collection of disconnected scripts. Every major methodological decision has a concrete operational expression, from file intake and transformation logic to selection policy and prediction reuse.",
        "Methodologically, this makes the project easier to defend in academic review. The report is not extrapolating from a vague implementation idea; it is documenting a workflow that can be inspected step by step, matched to source modules, and connected to saved experimental outputs."
    ])

    doc.add_page_break()
    add_heading(doc, "CHAPTER 4: RESULTS AND DISCUSSIONS", 1)
    add_heading(doc, "4.1 Experimental Setup", 2)
    add_long_paragraphs(doc, [
        "The workspace currently contains a saved classification run produced by the implemented NAS workflow. The environment is CPU-oriented and uses Flask, TensorFlow/Keras, Pandas, NumPy, and scikit-learn.",
        "The experiment compares multiple candidate architectures generated under bounded search conditions. Progress is streamed to the interface and summarized in training_report.json after completion.",
        f"The saved run includes {len(models)} candidate models. The average recorded training time across these candidates is {average_training_time:.2f} seconds, showing that the search remains lightweight enough for an interactive academic system.",
        "This setup matters when interpreting results. The project is not claiming large-scale benchmark superiority; rather, it demonstrates that a disciplined small-scale NAS pipeline can train, compare, store, and reuse multiple neural candidates under practical conditions.",
        "The experimental setup is also noteworthy because it combines algorithmic activity with interface-level visibility. The run is not only executed in the backend; it is also streamed outward through SSE so that the user can observe status changes and candidate progress. This helps bridge the usual gap between model training and user experience design.",
        "In academic terms, the setup is controlled rather than maximalist. It intentionally favours a small, interpretable candidate set over a large opaque search campaign. This makes the outcomes easier to explain, compare, and justify during report evaluation.",
        "The setup also exercises nearly every major subsystem in the application. A run touches file handling, task detection, preprocessing, candidate generation, model training, SSE reporting, artifact writing, and eventual reuse for prediction. Because of that breadth, the resulting evidence is suitable for both technical and academic discussion."
    ])
    add_heading(doc, "4.2 Dataset Description", 2)
    add_long_paragraphs(doc, [
        "The project is designed for mixed-type tabular medical datasets. The preprocessing logic assumes a final target column and feature columns before it.",
        "The literature notes repeatedly reference heart-disease prediction data, which serves as the motivating academic problem domain for the system.",
        "The preprocessing code is generic enough to support numeric and categorical fields together, which is important because many medical tables combine laboratory values with symbolic descriptors or yes/no indicators. This mixed-data assumption is visible in the separate imputation and transformation pipelines.",
        "Although the saved run report does not embed the full raw dataset, the project evidence strongly suggests a heart-disease-style binary classification use case. This interpretation is consistent with the uploaded file name, the literature survey sources, and the classification metrics stored in the generated report.",
        "This kind of dataset is a suitable target for the project's design because it is structured, moderate in scale, and likely to contain a mixture of physiological and categorical indicators. Such conditions reward careful preprocessing and measured model selection rather than excessive architectural complexity.",
        "The dataset discussion also highlights a practical limitation that the project handles reasonably well: the absence of embedded domain semantics in the code. The system does not hard-code disease-specific feature names, but instead relies on generic tabular processing. This increases reuse while still allowing domain-motivated interpretation in the report."
    ])
    add_heading(doc, "4.3 Candidate Models Generated", 2)
    model_rows = []
    for model in models:
        model_rows.append([
            str(model.get("candidate")),
            str(model.get("architecture")),
            str(model.get("total_params")),
            f"{model.get('training_time'):.2f}s" if isinstance(model.get("training_time"), (int, float)) else str(model.get("training_time")),
        ])
    add_report_table(doc, ["Candidate", "Architecture", "Parameters", "Training Time"], model_rows)
    add_long_paragraphs(doc, [
        "The candidate set illustrates how the NAS engine explores meaningful structural diversity without leaving the project's bounded design space. Some candidates use two hidden layers while others use three, and parameter counts vary from compact to moderately sized. This makes the comparison more informative than a single-model training run.",
        f"The smallest saved candidate is Candidate {smallest_model.get('candidate')} with {smallest_model.get('total_params')} parameters and architecture {format_architecture(smallest_model.get('architecture'))}. The fastest candidate is Candidate {fastest_model.get('candidate')} with a training time of {fastest_model.get('training_time'):.2f} seconds. These values show that the search process generates not only accuracy variation but also efficiency variation.",
        "Candidate diversity is academically valuable because it enables discussion of trade-offs. A stronger validation result may come from a compact network, a deeper network, or a different activation family. This moves the report beyond simple score reporting and toward reasoned model comparison.",
        "The generated candidates also show that the search space is neither trivial nor uncontrolled. Even within tight bounds, the project can produce architectures that differ in hidden-unit allocation, activation choice, parameter count, and training-time behaviour. This confirms that bounded search still leaves enough room for meaningful comparison.",
        "From a reporting standpoint, candidate diversity also enriches interpretation. It allows the project to discuss why some models are attractive because of efficiency, why others may be attractive because of test accuracy, and why the final selected model must balance those factors under a consistent ranking policy."
    ])
    add_heading(doc, "4.4 Training and Validation Performance", 2)
    perf_rows = []
    for model in models:
        perf_rows.append([
            str(model.get("candidate")),
            str(model.get("val_metric")),
            metric_percent(model.get("final_accuracy")) if model.get("final_accuracy") is not None else str(model.get("final_mse")),
            str(model.get("final_loss")),
        ])
    add_report_table(doc, ["Candidate", "Validation Metric", "Test Accuracy / MSE", "Test Loss"], perf_rows)
    add_long_paragraphs(doc, [
        "The performance table shows that candidate ranking cannot be reduced to a single intuitive guess. Several models achieve similar final test accuracy, yet their validation behaviour and complexity differ. This is precisely why a structured selection rule is necessary.",
        "Candidate 3 has the strongest validation metric at 88.89%, even though Candidates 1, 2, and 5 each show a slightly higher final test accuracy of 81.48%. Because the project selects the best model using validation performance, Candidate 3 becomes the preferred choice. This is methodologically sound and helps prevent test-set overfitting during model selection.",
        "Loss values add another layer of interpretation. Candidate 3 achieves the lowest stored final loss among the compared models, suggesting that its probability outputs may be better aligned with the true labels even when raw test accuracy is not uniquely maximal. This nuance is valuable in classification analysis because accuracy alone does not always capture the full quality of a model's predictions.",
        "The observed relationship between validation performance and final test accuracy is academically useful because it exposes the difference between model selection and model description. A candidate can have slightly lower final accuracy yet remain the correct validation-based choice. This is exactly the sort of detail that a serious machine learning report should discuss rather than conceal.",
        "The performance results also support the broader thesis of the project: architecture search is not about blindly maximizing size or depth. Instead, it is about finding a configuration whose observed behaviour across metrics is most credible under the adopted evaluation policy."
    ])
    add_report_table(doc, ["Observation", "Evidence from Current Run", "Interpretation"], [
        ["Best validation performer", f"Candidate {best.get('candidate', 'N/A')} at {metric_percent(best.get('val_metric'))}", "Selected for reuse and reporting"],
        ["Highest final test accuracy", f"Candidate {strongest_test_model.get('candidate') if strongest_test_model else 'N/A'}", "Shows that test score alone is not the selection rule"],
        ["Smallest architecture", f"Candidate {smallest_model.get('candidate')} with {smallest_model.get('total_params')} params", "Compactness remains competitive"],
        ["Fastest training", f"Candidate {fastest_model.get('candidate')} at {fastest_model.get('training_time'):.2f}s", "Runtime differences are visible even in bounded search"],
    ])
    add_long_paragraphs(doc, [
        "This summary table converts raw outcomes into explicit analytical findings. Instead of requiring the reader to infer why a particular number matters, the report explains how each observation contributes to the overall interpretation of the experiment.",
        "It also makes clear that the project evaluates models from multiple perspectives. A candidate may be attractive because it is fast, compact, or validation-strong, and the final selected model must be understood in relation to all three dimensions rather than through a single headline figure.",
        "This multi-perspective evaluation style is one of the strongest indicators that the report is grounded in serious machine learning reasoning. It does not hide ambiguities in the data; instead, it uses them as part of the explanation of why selection policy matters."
    ])
    add_report_table(doc, ["Candidate", "Validation Ranking Role", "Efficiency Role", "Interpretive Note"], [
        ["1", "Competitive but not best", "Moderate runtime, larger model", "Shows that higher test accuracy alone is insufficient"],
        ["2", "Competitive but not best", "Moderate runtime", "Demonstrates alternative relu-based configuration"],
        ["3", "Best by validation metric", "Compact and fast", "Chosen as best-model artifact"],
        ["4", "Lower validation score", "Fast compact baseline", "Useful contrast case in bounded search"],
        ["5", "Strong but second-tier validation", "Moderate runtime", "Illustrates trade-off between structure and ranking"],
    ])
    add_heading(doc, "4.5 Best Model Analysis", 2)
    add_bullets(doc, [
        f"Best candidate ID: {best.get('candidate', 'Not available')}",
        f"Architecture Details: {best.get('architecture', {})}",
        f"Validation Metric: {best.get('val_metric', 'Not available')}",
        f"Test Accuracy: {metric_percent(best.get('final_accuracy'))}",
        f"Training Behavior: {best.get('training_time', 'Not available')} seconds",
    ])
    add_long_paragraphs(doc, [
        "The saved best candidate in the current artifact is Candidate 3. It uses a tanh-based hidden architecture with two hidden layers and a compact parameter count.",
        "This result shows that a relatively small dense architecture can still perform competitively inside a constrained search space.",
        "Candidate 3 is especially interesting because it balances multiple desirable properties at once. It is not the largest model, not the slowest model, and not the most structurally complex model. Yet it achieves the best validation signal and the lowest recorded final loss. This is a strong example of why bounded NAS can be useful: it can discover efficient architectures that a user might not have chosen manually.",
        "The readable summary stored in the training report further highlights this interpretation by identifying Candidate 3 as both a fast-training and small-model solution. These descriptors reinforce the project's practical goal of finding credible models under limited-resource conditions rather than simply chasing the largest network possible.",
        "The best-model analysis also illustrates the value of combining quantitative and qualitative reasoning. Quantitatively, Candidate 3 leads on the validation metric and loss profile. Qualitatively, it is easy to justify as a reusable model because its compactness supports faster loading, simpler explanation, and more practical deployment in a lightweight system.",
        "This combination of compactness, validation strength, and preserved metadata is what makes Candidate 3 a convincing academic winner. The selection is not arbitrary and it is not hidden; it is explicit, evidenced, and reproducible from the saved workspace artifacts."
    ])
    if strongest_test_model:
        add_paragraph(doc, f"For comparison, the strongest final test accuracy in the current run belongs to Candidate {strongest_test_model.get('candidate')} at {metric_percent(strongest_test_model.get('final_accuracy'))}. The difference between this observation and the validation-based best-model choice demonstrates why the report treats evaluation as a multi-metric process rather than a one-number competition.")
    add_long_paragraphs(doc, [
        "The contrast between validation-based selection and test-based description is one of the most educational outcomes of the experiment. It shows why rigorous machine learning practice requires discipline in when and how metrics are used.",
        "This distinction is especially important for a system that persists and reuses the selected model. Once the project chooses a winner and saves it for later prediction, that decision should be methodologically defensible, not merely convenient."
    ])
    add_heading(doc, "4.6 User Interface and System Output", 2)
    add_long_paragraphs(doc, [
        "The training page supports upload, search-strategy selection, configuration values, and live updates. The prediction page renders input fields dynamically from the saved schema and displays task-aware outputs.",
        "The system output includes not only trained metrics but also reusable model files, preprocessing bundles, and run reports.",
        "From an academic software perspective, the interface is more than a cosmetic layer. It exposes the logic of the project to the user in a structured sequence: upload data, initiate search, monitor progress, review outputs, and perform prediction. That sequence helps evaluators see how the backend modules collaborate in practice.",
        "The output bundle produced by the system is also notable. Instead of returning only a numeric score, the workflow creates candidate-specific model files, a best-pipeline bundle, and a machine-readable report. This means the result of the experiment is reproducible, downloadable, and reusable for later demonstrations.",
        "The interface therefore performs a documentary role in addition to a control role. It becomes the visible surface through which the project's architecture, progress management, and reuse model can be demonstrated to a reader or examiner. That visibility is valuable because it links software design decisions to observable behaviour rather than leaving them buried only in backend code.",
        "Similarly, the downloadable outputs extend the life of the experiment beyond the moment of training. Reviewers can inspect the generated files, cross-check the report, and treat the project as a persistent system rather than a one-session demonstration.",
        "This interface-centered result is especially relevant in an academic project because demonstration quality matters alongside technical correctness. A system that can be understood, navigated, and reused through clear pages and outputs is easier to evaluate and more valuable for future student cohorts."
    ])
    add_heading(doc, "4.7 Discussion of Results", 2)
    add_long_paragraphs(doc, [
        "The results demonstrate that bounded neural architecture search can be applied meaningfully to tabular medical prediction without requiring large-scale infrastructure.",
        "A major strength of the project is its full integration of preprocessing, training, monitoring, storage, and prediction inside one workflow.",
        "The current run does not claim to surpass the strongest manually engineered medical prediction systems reported in the literature. Instead, its significance lies in showing that automated architecture comparison can be embedded in a user-facing application while preserving methodological discipline. In that sense, the project's main contribution is systems integration grounded in legitimate machine learning practice.",
        "The results also support a practical lesson: for tabular structured data, compact networks can be highly competitive. This aligns with broader observations that massive depth is not always necessary for table-based prediction tasks, especially when preprocessing is strong and the dataset is moderate in size.",
        "Another discussion point concerns the choice of search metric. Because the project uses validation performance for model selection, the reported best model is stable with respect to the stated methodology. This is preferable to choosing a candidate based only on final test accuracy after viewing all results, which would weaken the credibility of the evaluation.",
        "The discussion also highlights the importance of software embodiment. Many projects describe a promising machine learning idea but stop short of integrating it into a maintainable and demonstrable application. Here, the results are inseparable from the system that produced them: the backend, preprocessing pipeline, search engine, persistence logic, and prediction interface all contribute to the final value of the project.",
        "Another interpretive point is that the experiment reveals operational trade-offs as well as predictive ones. Candidate training times differ, parameter counts differ, and selection outcomes differ. This means the project can discuss efficiency, reproducibility, and practical deployability in addition to predictive quality, making the final chapter more balanced than a pure score-based comparison.",
        "The overall discussion therefore supports the claim that the project is academically meaningful not because it is the largest or most complex NAS implementation, but because it demonstrates the complete reasoning chain from input management to architecture comparison to artifact reuse in a coherent, inspectable form.",
        "The project therefore contributes a form of practical evidence that is often missing from high-level discussions of AutoML and NAS. It shows what happens when search is embedded in a web-facing, artifact-preserving workflow rather than presented only as an isolated optimization procedure.",
        "In that sense, the project offers value not only through its model outcomes but also through its explanation of process. The report can discuss candidate diversity, ranking discipline, artifact persistence, and interface behaviour together, which is uncommon in short prototype reports and strengthens the academic character of the submission.",
        "One of the deepest lessons from the current results is that small-scale AutoML systems can still teach large-scale ideas. The project shows how search spaces are defined, why search strategy matters, how validation must be separated from test reporting, how artifact preservation supports reproducibility, and why deployment concerns appear even in a prototype. These are not minor details; they are central themes of modern applied machine learning. By expressing them through a constrained but functioning application, the project converts high-level concepts into something concrete and inspectable.",
        "The results also emphasize the importance of system coherence. A strong metric alone would not make the project convincing if preprocessing were inconsistent, the best model were not saved properly, or the prediction route could not reuse the stored artifacts. What makes the current outputs credible is that they emerge from a workflow where each stage supports the next. Upload conventions enable preparation, preparation enables meaningful search, search produces traceable outcomes, and those outcomes feed prediction and documentation. The value of the results therefore lies as much in this continuity as in any single number in the report.",
        "Another insight from the experiment is that bounded search is not merely a compromise forced by resource limits. It can also be a design advantage. Because the project explores a restricted but meaningful family of dense architectures, the reader can interpret the results without being overwhelmed by excessive architectural variation. This interpretability matters in academic reporting, where the goal is not only to achieve performance but also to explain why the observed behaviour is reasonable. The current project succeeds precisely because it makes architecture search understandable rather than mysterious.",
        "The results chapter further demonstrates that project quality should be judged by evidence discipline as well as implementation effort. The software records candidate-level outputs, preserves the winning architecture, and exposes a readable summary of why that model was considered best. This evidential discipline makes the report more robust. It also prepares the project for future extension, because later comparisons can build on preserved candidate histories instead of starting from undocumented experimentation. In this sense, the current results are not just the end of one run; they are the beginning of a reusable evaluation tradition within the workspace.",
        "Finally, the discussion confirms that the project's main achievement is not a claim of universal superiority over all heart disease prediction systems or all NAS methods. Its real achievement is disciplined integration. The application demonstrates how a modern machine learning idea can be packaged into a modular backend, a dynamic interface, a reproducible artifact set, and an academically defensible report. That integrated success gives the project enduring value even if future work later improves the search space, the metrics, or the presentation layer."
    ])
    add_report_table(doc, ["Discussion Theme", "Evidence in Project", "Why It Matters"], [
        ["Methodological discipline", "Validation-based best-model choice", "Prevents careless use of the test set"],
        ["Operational transparency", "SSE monitoring and stored artifacts", "Makes the workflow inspectable"],
        ["Efficient experimentation", "Compact models and fast runtimes", "Shows feasibility on limited resources"],
        ["Reusability", "Prediction module and saved pipeline", "Turns training output into a working service"],
        ["Educational value", "Multiple search modes in one app", "Supports comparison and learning"],
    ])
    add_long_paragraphs(doc, [
        "A deeper reading of the results suggests that the project succeeds as an evidence-producing system. Each candidate is more than a temporary training attempt; it becomes a documented experimental unit with architecture details, runtime measurements, and evaluation metrics. This means the project supports reflective analysis after training rather than forcing the reader to rely on memory or informal observation. In academic documentation, that shift from transient activity to preserved evidence is highly significant because it turns the experiment into a repeatable and discussable object.",
        "The results chapter also reveals how software constraints can shape meaningful methodological choices. Because the system is intended to remain lightweight and interactive, the search space and training regime are intentionally constrained. Yet those constraints do not make the experiment trivial. Instead, they create a setting where trade-offs become visible: some candidates are faster, some are larger, some score better on validation, and some look appealing by test accuracy alone. This kind of constrained variation is often more instructive for learning than a very large opaque search whose internal logic is hard to explain.",
        "Another important result concerns the integration of monitoring and artifact generation. The application does not wait until the end of the run to become informative. It streams progress during execution and then stores outputs afterward. This dual visibility—live observation followed by persistent evidence—strengthens both the user experience and the credibility of the documentation. Reviewers can treat the project as a system that communicates throughout the experimental lifecycle rather than as a black box that reveals only a final score.",
        "The candidate comparison further demonstrates that model quality in this context is multidimensional. If the report considered only final test accuracy, the interpretation would be incomplete and potentially misleading. If it considered only runtime, it would undervalue predictive behaviour. If it considered only compactness, it would miss the methodological reason for validation-based selection. By retaining all of these dimensions together, the project enables a more balanced academic discussion of what it means for one candidate to be preferable to another.",
        "The saved best model also has value beyond its own metric values. It proves that the system can complete the entire path from data upload through preprocessing, architecture search, ranking, persistence, and prediction reuse. In many student projects, the workflow breaks at one of these transitions: preprocessing is not reusable, model metadata is lost, or inference cannot be reproduced reliably. The present result suggests that these transitions have been handled with enough care to support a believable end-to-end application narrative.",
        "Finally, the results have educational value precisely because they are modest and interpretable. The project does not rely on spectacular benchmark claims to justify itself. Instead, it shows that a bounded NAS workflow can be implemented responsibly, that its outputs can be stored and explained, and that its best model can be defended through a clear selection rationale. This kind of measured success is often more useful in an academic project report than exaggerated claims that are difficult to verify or reproduce."
    ])
    add_heading(doc, "4.8 Limitations of the System", 2)
    add_bullets(doc, [
        "The project is limited to tabular dense-network search.",
        "The search space is intentionally narrow for practical runtime reasons.",
        "The system is an academic prototype, not a regulated clinical solution.",
        "Figures are listed textually rather than embedded as chart images.",
    ])
    add_long_paragraphs(doc, [
        "These limitations should be interpreted carefully. A narrow search space reduces the chance of discovering unconventional high-performing structures, yet it also keeps the system explainable and fast enough for classroom use. Similarly, the absence of hospital-grade controls does not reduce the project's value as a prototype; it simply defines the maturity level of the software.",
        "Another limitation is the dependence on the last-column target convention. While practical, this assumption requires the dataset to be prepared in a conventional tabular form. A future production-oriented version would likely include explicit target selection and stronger validation around schema semantics.",
        "Finally, evaluation remains tied to the available saved run artifacts. More extensive experiments across multiple medical datasets, repeated random seeds, and confusion-matrix-style diagnostics would strengthen the empirical claims that future versions of the report can make.",
        "A further limitation is that the current visualization emphasis is textual and tabular rather than graph-heavy. While this is acceptable for a document-focused academic submission, richer visual analytics would improve accessibility for readers who prefer curve-based or dashboard-based interpretation of candidate behaviour.",
        "There is also a scope limitation in user interaction design. The system supports a clear training and prediction workflow, but it does not yet provide deeper experiment-management capabilities such as comparing historical runs, annotating experiments, or switching between multiple saved best pipelines from the interface."
    ])

    doc.add_page_break()
    add_heading(doc, "CHAPTER 5: CONCLUSION AND FUTURE SCOPE", 1)
    add_heading(doc, "5.1 Conclusion", 2)
    add_long_paragraphs(doc, [
        f"{PROJECT_TITLE} demonstrates that neural architecture search for tabular medical prediction can be implemented as a practical software system. The project unifies data preparation, candidate generation, training, evaluation, artifact persistence, and prediction support in one workflow.",
        "The project is especially valuable in academic settings because it connects theoretical NAS ideas to a usable end-to-end application.",
        "The completed implementation shows that automation in model design need not be isolated from software engineering discipline. By preserving preprocessing, storing candidate artifacts, streaming progress, and reusing the chosen model for prediction, the system demonstrates a full lifecycle from dataset upload to inference.",
        "The project therefore succeeds on two levels. At the machine learning level it produces and evaluates alternative neural architectures for a tabular classification task. At the system level it packages that capability inside a coherent Flask application whose behaviour can be observed, documented, and reused.",
        "The conclusion that emerges from the completed work is not merely that NAS can function on tabular data. It is that a small but thoughtfully engineered NAS application can serve as a bridge between machine learning methodology, software engineering practice, and academic documentation requirements. This integrated success is the defining contribution of the project.",
        "The saved artifacts and report evidence also confirm that the project produces inspectable outcomes rather than abstract promises. The selected model, the preserved preprocessing bundle, the candidate files, and the final training report collectively show that the system moved from design intent to measurable execution.",
        "The strongest final conclusion is therefore one of successful integration. The project does not leave its main ideas scattered across disconnected artifacts. It brings them together into a coherent platform where the technical workflow, the interface workflow, and the documentation workflow all reinforce one another.",
        "This makes the project particularly suitable for academic evaluation because it can be defended from several angles at once: algorithm design, preprocessing rigor, software architecture, interface clarity, and artifact-based reproducibility. Few small projects make all of those dimensions visible in one coherent submission."
    ])
    add_heading(doc, "5.2 Achievements of the Project", 2)
    add_report_table(doc, ["Achievement Area", "Description"], [
        ["Data Pipeline", "Implemented preprocessing, task detection, and feature schema generation"],
        ["Model Search", "Implemented random, evolutionary, and progressive NAS flows"],
        ["System Integration", "Integrated Flask UI, SSE updates, downloads, and prediction support"],
        ["Documentation", "Prepared academic and technical documentation aligned with the real implementation"],
    ])
    add_long_paragraphs(doc, [
        "These achievements are not isolated checklist items; they form a connected software ecosystem. The data pipeline feeds the NAS engine, the NAS engine feeds the reporting layer, and the reporting layer feeds both academic documentation and later prediction workflows.",
        "A significant achievement is that the project preserves verifiable outputs. The saved report, the best pipeline bundle, and the candidate model files provide objective evidence that the workflow executed and that the best-model decision can be traced to specific artifacts.",
        "Another achievement is conceptual clarity. The project makes it possible to explain each stage of the workflow in a way that maps directly to code and outputs. This clarity is especially useful in academic assessment, where a strong project is judged not only by what it does but also by how transparently it can be justified.",
        "The project also achieves a useful compromise between ambition and feasibility. It adopts the important idea of automated architecture search without depending on expensive training infrastructure or overly complex experimental protocols. This balance is one of the reasons the implementation is well suited to a student project environment.",
        "These achievements also indicate that the project can serve as a starting point for later teams. Because the architecture, saved outputs, and documentation are aligned, subsequent development can build on a stable base instead of reverse-engineering undocumented decisions."
    ])
    add_heading(doc, "5.3 Limitations", 2)
    add_long_paragraphs(doc, [
        "The project does not currently include advanced explainability features, GPU cluster support, user authentication, or deployment-grade medical compliance controls.",
        "It also does not yet provide broad experiment management features such as multi-run comparison dashboards, persistent user accounts, audit logging, or configurable cross-validation studies. These omissions are acceptable for a focused academic prototype but should be recognized when positioning the system against industrial AutoML platforms.",
        "Another limitation is that the present report is grounded in the artifacts of the current workspace run. While this makes the document accurate and reproducible, it also means broader empirical generalization requires more saved experiments across additional datasets and parameter settings.",
        "There is also a limitation in how operational analytics are presented to the reader. The system preserves essential metrics, but it does not yet offer a wider set of automatically generated analytical summaries such as confusion matrices, class-wise precision and recall, or probability calibration curves for classification tasks.",
        "Finally, the project remains a prototype with a single-user execution mindset. Concurrency safeguards exist, but the broader concerns of multi-user deployment, quota management, secure model isolation, and auditability would need further engineering in a production-oriented continuation.",
        "These limitations do not weaken the project's academic value; instead, they help locate it accurately within the prototype-to-product spectrum. A strong report should not hide what the system does not yet solve. By stating these constraints openly, the project strengthens its credibility and prepares a clear agenda for future improvement."
    ])
    add_heading(doc, "5.4 Future Scope", 2)
    add_bullets(doc, [
        "Expand the search space with more regularization and topology options.",
        "Add explainability and calibration outputs.",
        "Introduce user authentication and run-history management.",
        "Add richer visualization and automated testing.",
        "Extend support for more configurable NAS parameters from the UI.",
    ])
    add_long_paragraphs(doc, [
        "Future work can deepen both the machine learning and the software dimensions of the project. On the search side, dropout, batch normalization, residual-style tabular blocks, and wider activation choices could provide a richer architecture space. On the evaluation side, confusion matrices, ROC analysis, calibration checks, and repeated trials would improve scientific depth.",
        "From a systems perspective, future scope includes user authentication, saved experiment histories, administrative dashboards, and richer visualization of candidate training curves. Those extensions would make the application more suitable for repeated departmental use.",
        "Another promising direction is the integration of explainability. Because medical prediction applications benefit from transparency, a later version could expose feature attributions, local explanation summaries, or confidence diagnostics together with the final prediction output.",
        "Finally, the project can evolve toward a more configurable educational AutoML laboratory. By allowing users to control additional search parameters, compare strategies side by side, and export richer reports, the same foundation could support deeper coursework, student mini-projects, and future research exploration.",
        "A particularly valuable future direction would be the addition of experiment reproducibility controls such as fixed-seed run tracking, side-by-side run comparison, and archived metadata snapshots. These features would strengthen the project as both a teaching tool and a small-scale research support platform.",
        "Another future path is closer integration with documentation generation itself. Since the project already saves structured reports and artifacts, a later version could automatically generate richer evidence sections, candidate comparison charts, and update-ready academic summaries from each completed run.",
        "Over time, the same system could become a reusable departmental asset: a platform where students upload datasets, compare search strategies, inspect artifacts, and produce consistent documentation. This long-term educational potential gives the current prototype value beyond its immediate implementation scope.",
        "Future scope also includes stronger deployment realism, such as containerized execution, environment validation, and role-based access to saved experiments. These features would not change the conceptual heart of the project, but they would extend its utility in collaborative settings.",
        "A parallel future direction is richer dataset governance. Allowing explicit target-column selection, dataset preview checks, and configurable preprocessing strategies would make the system more flexible while preserving the same core commitment to reproducibility and traceable artifact generation.",
        "Taken together, these future directions suggest that the current implementation is not an endpoint but a stable foundation. It already embodies the core workflow, and later work can enrich it horizontally through better analytics and vertically through stronger deployment practices without discarding the present architecture.",
        "Future scope should also include richer comparative reporting. Since the project already stores candidate-level information, later versions could automatically generate side-by-side architectural summaries, trend reports across runs, and more advanced evidence sections for academic documentation. This would reduce manual report preparation effort while strengthening the interpretability of each experiment.",
        "Another path forward is stronger pedagogical instrumentation. The application could include explanatory overlays, step-wise workflow notes, or strategy comparison guides that help students understand why random, evolutionary, and progressive search behave differently. Such features would make the tool even more effective as a teaching platform rather than only a training interface.",
        "From a research-support perspective, the system could evolve into a lightweight experiment lab for structured data. Multiple datasets, repeated seeds, user-defined search budgets, and archived result comparisons would allow the same software foundation to support semester projects, mini-research studies, and internal benchmarking exercises.",
        "A clinically oriented continuation could also add stronger governance around inputs and outputs. While the current project is explicitly academic, future versions might incorporate structured validation of sensitive fields, clearer prediction disclaimers, confidence reporting standards, and audit-ready artifact records. These additions would improve the platform's maturity without changing its fundamental design logic.",
        "The most important future insight is that the project has already solved the hardest integration problem: connecting upload, preprocessing, search, evaluation, storage, and prediction into one coherent workflow. Because that foundation exists, later extensions can focus on depth rather than rebuilding the system from scratch. This gives the project strong long-term value as a platform for continued development.",
        "In other words, the current system should be viewed as a capable baseline platform. Its immediate contribution is a complete academic prototype, but its longer-term contribution is the availability of a reusable structure on which richer analysis, deployment, and teaching features can be layered systematically.",
        "A valuable next step would be to make experimentation itself a first-class object in the interface. Rather than treating each run as a separate event that is remembered informally, future versions could preserve named experiment sessions, configuration snapshots, metric summaries, and candidate comparisons across time. This would make the software not only a training tool but also a lightweight experiment-management environment suitable for classroom and mini-project usage.",
        "Future scope can also deepen methodological rigor by supporting more flexible evaluation procedures. For example, repeated train-validation-test cycles, user-selectable random seeds, and cross-validation-inspired summary reporting would help quantify the stability of the selected model. These changes would be especially useful if the platform is later applied to multiple medical datasets, where robustness across runs matters as much as peak performance in any single run.",
        "The prediction module offers another rich avenue for extension. At present it successfully reuses the stored pipeline and model, which is already a major achievement for a prototype. Future versions could build on this by adding richer output narratives, probability calibration displays, confidence threshold options, and clearer input guidance derived from the feature schema. Such improvements would make the system more informative for end users while remaining faithful to the current architectural pattern.",
        "From a software engineering perspective, the project could also evolve toward a more maintainable product ecosystem. Automated tests around preprocessing, report generation, route behavior, and prediction consistency would make future modifications safer. Versioned artifact management could ensure that older reports and models remain accessible even after new runs are added. These changes would strengthen reliability without requiring a fundamental redesign.",
        "Longer term, the platform could become an institutional teaching asset. Faculty could use it to demonstrate the full lifecycle of a structured-data ML project, students could compare search strategies in a controlled environment, and project teams could generate consistent documentation backed by real artifacts. In that scenario, the value of the current prototype extends beyond its own immediate results: it becomes the seed of a reusable educational and experimental infrastructure."
    ])
    add_report_table(doc, ["Future Area", "Example Enhancement", "Expected Benefit"], [
        ["Search space", "Add dropout and normalization options", "Richer architecture exploration"],
        ["Evaluation", "Add confusion matrix and ROC summaries", "Stronger empirical interpretation"],
        ["Interface", "Run history and experiment comparison", "Improved usability and reflection"],
        ["Deployment", "Containerized execution and access control", "Better collaborative demonstration support"],
        ["Documentation", "Auto-generated evidence sections", "Faster academic reporting with better traceability"],
    ])
    add_long_paragraphs(doc, [
        "The future-oriented view of the project is therefore both practical and strategic. Practically, the next improvements are clear: richer evaluation, better experiment management, stronger explainability, and more structured deployment support. Strategically, the project already possesses the most important asset required for those improvements—a coherent baseline architecture in which data preparation, model search, reporting, and prediction are already connected. Because that foundation is in place, the path to a stronger next version is incremental rather than disruptive.",
        "This final point matters in an academic context because it shows that the project has continuity. It is not a one-time exercise whose value ends with the current report. Instead, it is a reusable framework that can support future student work, more rigorous experiments, and richer documentation. That continuity gives the current implementation a stronger educational significance than a prototype that solves only a narrow isolated task.",
        "A final strategic observation is that the project already demonstrates the rare combination of machine learning workflow completeness and documentation readiness. Many prototypes can train a model, and many reports can describe a workflow in abstract terms, but fewer systems can do both in a coordinated manner. Here, the workflow produces saved models, a saved preprocessing bundle, a machine-readable report, and a user-facing prediction interface, all of which reinforce the written academic narrative. This makes the project more durable as a departmental reference point because future contributors can inherit not just code, but also a working pattern for how code, artifacts, and documentation should relate.",
        "For that reason, the long-term importance of the current implementation should not be judged only by the numerical performance of a single run. Its broader contribution is infrastructural. It shows how a bounded NAS system for tabular medical prediction can be organized so that experiments are observable, outputs are preserved, and conclusions remain defensible. Future teams may improve metrics, broaden datasets, or modernize the interface, but they will be doing so on top of an already coherent foundation. In academic project work, that kind of reusable structure is itself a significant achievement because it turns one successful submission into a platform for later learning and experimentation."
    ])

    doc.add_page_break()
    add_heading(doc, "REFERENCES", 1)
    for paper in papers:
        add_paragraph(
            doc,
            f"{paper.get('authors', 'Authors not specified')} ({paper.get('year', 'n.d.')}). {paper.get('title', 'Untitled')}. {paper.get('source', 'Source not specified')}."
        )

    doc.add_page_break()
    add_heading(doc, "LIST OF PUBLICATIONS", 1)
    add_paragraph(doc, "At the time of report preparation, no separate journal or conference publication record is present in the current workspace artifacts for this project.")

    output = DOC_DIR / "1_PROJECT_REPORT.docx"
    output.unlink(missing_ok=True)
    doc.save(output)


def generate_srs() -> None:
    doc = new_document("Software Requirements Specification (SRS)", PROJECT_TITLE)
    add_heading(doc, "1. Introduction", 1)
    add_heading(doc, "1.1 Purpose", 2)
    add_long_paragraphs(doc, [
        "This Software Requirements Specification defines the functional and non-functional requirements of the implemented NAS Medical Model Generator system. The document is based strictly on the actual Python, Flask, preprocessing, neural training, and interface code present in the workspace.",
        "The purpose of the system is to provide a web-based workflow for tabular dataset upload, automatic preprocessing, dense neural network architecture search, model training, model comparison, best-model selection, artifact saving, and prediction. The system is designed for academic and prototype use rather than production clinical deployment."
    ])
    add_heading(doc, "1.2 Scope", 2)
    add_long_paragraphs(doc, [
        "The software accepts CSV and Excel datasets, detects whether the task is classification or regression, prepares the dataset using missing value handling, encoding, and normalization, and splits the data into train, validation, and test partitions using an effective 60/20/20 ratio.",
        "The system then applies one of three implemented search strategies—random search, evolutionary search, or progressive search—to generate and train dense neural network candidates. It evaluates candidates using accuracy and loss for classification or MSE and loss for regression, selects the best model using validation performance, saves the resulting artifacts, and exposes a prediction workflow through a dynamic form generated from saved feature schema metadata."
    ])
    add_heading(doc, "1.3 Definitions", 2)
    add_bullets(doc, [
        "NAS: Neural Architecture Search used in this project to compare dense neural network candidates.",
        "SSE: Server-Sent Events used by the `/stream` endpoint for real-time training updates.",
        "Feature Schema: Metadata describing input fields, including numeric ranges and categorical choices, used to build the prediction form dynamically.",
        "Best Model: The candidate selected using validation performance after candidate training and evaluation.",
        "Artifact: A saved output file such as a `.keras` model, `.joblib` preprocessing bundle, or `training_report.json` file.",
    ])

    add_heading(doc, "2. Overall Description", 1)
    add_heading(doc, "2.1 Product Perspective", 2)
    add_long_paragraphs(doc, [
        "The product is a client-server machine learning application built with Flask on the backend and HTML templates on the frontend. It integrates three major implementation areas: a web interface for user interaction, a preprocessing/data pipeline for tabular datasets, and a NAS engine for dense neural network training and comparison.",
        "The system is not a general deep learning platform. It is limited to dense feed-forward neural networks with 1 to 4 hidden layers, units selected from 16, 32, 64, and 128, hidden activations limited to ReLU and Tanh, and Adam optimizer with learning rate 0.001."
    ])
    add_heading(doc, "2.2 Product Functions", 2)
    add_bullets(doc, [
        "Upload CSV, XLS, and XLSX datasets through the training interface.",
        "Validate file type and reject unsupported uploads.",
        "Automatically detect classification or regression from the target column.",
        "Handle missing values and perform encoding plus normalization for tabular input.",
        "Split data into train, validation, and test sets using a 60/20/20 workflow.",
        "Generate dense neural network candidates using random, evolutionary, or progressive search.",
        "Train candidates and stream real-time updates using SSE.",
        "Evaluate candidates using accuracy/loss or MSE/loss according to task type.",
        "Provide a model comparison API and best-model summary.",
        "Save candidate models, best pipeline artifacts, and training report files.",
        "Generate a dynamic prediction form from saved feature schema metadata.",
        "Run inference using saved preprocessing and best-model artifacts.",
    ])
    add_heading(doc, "2.3 User Classes", 2)
    add_bullets(doc, [
        "Student User: uploads a dataset, starts a NAS run, monitors progress, views comparison results, and uses the prediction interface.",
        "Project Evaluator or Faculty Reviewer: inspects the generated outputs, model comparison data, and saved artifacts for academic evaluation.",
        "Developer or Maintainer: reviews logs, source modules, and saved artifacts to improve or validate the system behavior.",
    ])
    add_heading(doc, "2.4 Operating Environment", 2)
    add_bullets(doc, [
        "Python-based execution environment.",
        "Flask web server for route handling and page rendering.",
        "TensorFlow/Keras for dense neural network building and training.",
        "Pandas, NumPy, scikit-learn, and joblib for dataset preparation and artifact persistence.",
        "CPU-oriented execution environment as configured in the application startup.",
        "Browser-based frontend using the implemented HTML templates.",
    ])
    add_heading(doc, "2.5 Constraints", 2)
    add_bullets(doc, [
        "Only dense neural networks are supported by the implemented search engine.",
        "The search space is limited to 1–4 hidden layers.",
        "Hidden-layer units are limited to 16, 32, 64, and 128.",
        "Hidden activations are limited to ReLU and Tanh.",
        "Optimizer is Adam with learning rate 0.001.",
        "Only one training run is allowed at a time.",
        "User-provided candidate count, batch size, and epoch values are clamped to safe ranges by the backend.",
        "The last dataset column is treated as the target variable by convention.",
    ])

    add_heading(doc, "3. Functional Requirements", 1)
    functional_sections = {
        "3.1 Dataset Upload": [
            "The system shall accept CSV, XLS, and XLSX files through the upload route.",
            "The system shall reject missing files and unsupported file extensions.",
            "The system shall store uploaded files using sanitized timestamped filenames.",
            "The system shall reject a new upload request if another training run is already in progress.",
        ],
        "3.2 Data Preprocessing": [
            "The system shall read CSV and Excel files into a tabular DataFrame.",
            "The system shall drop rows with missing target values.",
            "The system shall identify numeric and categorical feature columns.",
            "The system shall apply median imputation and StandardScaler to numeric columns.",
            "The system shall apply most-frequent imputation and OneHotEncoder to categorical columns.",
            "The system shall fit preprocessing on training data and reuse the same fitted transformation for validation, test, and prediction paths.",
        ],
        "3.3 Task Detection": [
            "The system shall detect classification when the target is non-numeric.",
            "The system shall detect classification when the target is numeric with limited unique values according to implemented heuristic logic.",
            "The system shall detect regression when the target is numeric with many unique values.",
        ],
        "3.4 NAS Model Generation": [
            "The system shall support random, evolutionary, and progressive search strategies.",
            "The system shall generate dense architectures with 1–4 hidden layers only.",
            "The system shall use unit counts from 16, 32, 64, and 128 only.",
            "The system shall use ReLU or Tanh as hidden activation functions only.",
            "The system shall skip candidates whose parameter count exceeds the implemented maximum limit.",
        ],
        "3.5 Model Training": [
            "The system shall compile classification models using Adam optimizer, categorical cross-entropy loss, and accuracy metric.",
            "The system shall compile regression models using Adam optimizer, mean squared error loss, and MSE metric.",
            "The system shall train candidates using TensorFlow/Keras and tf.data pipelines.",
            "The system shall apply early stopping during candidate training.",
        ],
        "3.6 Model Evaluation": [
            "The system shall evaluate candidates on validation and test data after training.",
            "The system shall record accuracy and loss for classification tasks.",
            "The system shall record MSE and loss for regression tasks.",
            "The system shall preserve candidate-level results in the generated report data.",
        ],
        "3.7 Best Model Selection": [
            "The system shall select the best candidate using validation performance.",
            "For classification, the system shall prefer higher validation accuracy.",
            "For regression, the system shall prefer lower validation MSE through the implemented ranking logic.",
            "The system shall expose the selected best model in the model comparison output and prediction workflow.",
        ],
        "3.8 Model Storage": [
            "The system shall save each trained candidate as a `.keras` model file.",
            "The system shall save the fitted preprocessor, feature columns, feature schema, label encoder when applicable, task type, and target name as a `.joblib` bundle.",
            "The system shall save a `training_report.json` file containing task type, all model summaries, best model information, and readable summary text.",
            "The system shall support downloading the best model, training report, and individual candidate models through implemented routes.",
        ],
        "3.9 Prediction Module": [
            "The system shall generate a feature schema for the prediction interface from training data.",
            "The system shall expose feature schema metadata through the prediction schema route.",
            "The system shall build a single-row prediction input using saved feature order and types.",
            "The system shall reuse the saved preprocessing pipeline before running inference.",
            "The system shall return predicted class label and probabilities for classification or predicted numeric value for regression.",
        ],
        "3.10 Real-Time Training Visualization (SSE)": [
            "The system shall stream status, model information, epoch-end updates, and result events through Server-Sent Events.",
            "The system shall keep the web application responsive during training by running the training process in a background thread.",
        ],
        "3.11 Model Comparison": [
            "The system shall expose a model comparison API that returns task type, best model, all model summaries, and a readable summary.",
            "The system shall support UI consumption of candidate comparison data after report generation.",
        ],
    }
    for title, items in functional_sections.items():
        add_heading(doc, title, 2)
        add_bullets(doc, items)

    add_heading(doc, "4. Non-Functional Requirements", 1)
    add_heading(doc, "4.1 Performance", 2)
    add_bullets(doc, [
        "The system shall keep the web interface responsive during model training by delegating training to a background thread.",
        "The system shall stream incremental training updates using SSE instead of waiting for run completion.",
        "The system shall limit candidate count, batch size, epochs, and model size using implemented safeguards.",
    ])
    add_heading(doc, "4.2 Usability", 2)
    add_bullets(doc, [
        "The system shall provide separate training and prediction pages.",
        "The system shall support common spreadsheet-style dataset formats.",
        "The system shall create prediction inputs dynamically from saved feature schema metadata.",
        "The system shall provide downloadable outputs for inspection and reuse.",
    ])
    add_heading(doc, "4.3 Reliability", 2)
    add_bullets(doc, [
        "The system shall reject unsupported uploads and invalid NAS types.",
        "The system shall prevent overlapping training runs.",
        "The system shall preserve trained artifacts and report files after successful execution.",
        "The system shall reuse the same fitted preprocessor during prediction to reduce train-predict inconsistency.",
    ])
    add_heading(doc, "4.4 Maintainability", 2)
    add_bullets(doc, [
        "The system shall remain modular through separate application, data pipeline, and NAS engine modules.",
        "The system shall keep runtime state in a dedicated structured container.",
        "The system shall keep model comparison, preprocessing, and prediction logic separated by responsibility.",
    ])
    add_heading(doc, "4.5 Security", 2)
    add_bullets(doc, [
        "The system shall sanitize uploaded filenames before saving them.",
        "The system shall restrict accepted dataset formats to CSV and Excel files.",
        "The system shall avoid arbitrary prediction schema creation by relying on saved training artifacts.",
        "The system shall expose only the implemented academic prototype routes and artifact downloads.",
    ])

    output = DOC_DIR / "2_SRS_DOCUMENT.docx"
    output.unlink(missing_ok=True)
    doc.save(output)


def generate_sdd() -> None:
    doc = new_document("Software Design Document (SDD)", PROJECT_TITLE)
    add_heading(doc, "1. Introduction", 1)
    add_long_paragraphs(doc, [
        "This Software Design Document describes the design of the implemented NAS Medical Model Generator system. The document is written from the actual workspace implementation and focuses on the deployed client-server structure, preprocessing flow, dense-network search engine, model persistence, and prediction workflow.",
        "The design targets tabular machine learning problems only. The implemented system does not use CNN, RNN, GAN, or other advanced architectures. Instead, it uses dense neural networks generated within a limited search space and managed through a Flask-based web application."
    ])

    add_heading(doc, "2. System Architecture", 1)
    add_heading(doc, "2.1 Client-Server Architecture", 2)
    add_long_paragraphs(doc, [
        "The system follows a client-server pattern. The client side consists of browser-accessible HTML pages for training and prediction. The server side is implemented with Flask and coordinates upload handling, background training, streaming updates, artifact access, feature schema delivery, and prediction requests.",
        "The design separates user interaction from machine learning execution. User requests are handled through Flask routes, while training is delegated to a background worker thread so that the interface remains responsive during model search and candidate training."
    ])
    add_heading(doc, "2.2 Frontend, Backend, and ML Engine", 2)
    add_bullets(doc, [
        "Frontend Layer: training page and prediction page rendered from HTML templates.",
        "Backend Layer: Flask routes, runtime state container, locks, artifact download handlers, and SSE streaming endpoint.",
        "ML Engine Layer: dataset preparation, feature schema generation, dense architecture generation, candidate training, evaluation, and model selection.",
    ])
    add_report_table(doc, ["Layer", "Implemented Elements", "Design Role"], [
        ["Frontend", "`templates/index.html`, `templates/predict.html`", "Collect input, display progress, show prediction output"],
        ["Backend", "`app.py` routes, runtime state, event queue, model cache", "Manage requests, concurrency, SSE, downloads, and inference"],
        ["ML Engine", "`data_pipeline.py`, `nas_engine.py`", "Prepare data, search dense architectures, train, evaluate, and save results"],
    ])

    add_heading(doc, "3. Module Design", 1)
    add_heading(doc, "3.1 Data Pipeline Module", 2)
    add_long_paragraphs(doc, [
        "The Data Pipeline Module is implemented in `data_pipeline.py`. It reads CSV or Excel data, validates the input structure, drops rows with missing target values, separates features from the target, identifies numeric and categorical columns, and builds preprocessing pipelines using scikit-learn components.",
        "Numeric data is processed with median imputation and StandardScaler. Categorical data is processed with most-frequent imputation and OneHotEncoder. The module also performs the effective 60/20/20 train-validation-test split and packages the processed arrays and metadata into the `PreparedData` structure."
    ])
    add_heading(doc, "3.2 Feature Schema Module", 2)
    add_long_paragraphs(doc, [
        "Feature schema generation is implemented as part of the data pipeline. It creates metadata for every feature column, including type information and either numeric hints or categorical choices.",
        "This module supplies the data required by the prediction interface to build input fields dynamically. It also ensures that the prediction path stays aligned with the structure used during model training."
    ])
    add_heading(doc, "3.3 NAS Engine", 2)
    add_long_paragraphs(doc, [
        "The NAS Engine is implemented in `nas_engine.py`. It generates dense neural network candidates using three supported strategies: random search, evolutionary search, and progressive search.",
        "The engine restricts search to dense neural networks with 1 to 4 hidden layers, units selected from 16, 32, 64, and 128, and hidden activations limited to ReLU and Tanh. It also enforces the maximum parameter guard defined in the implementation."
    ])
    add_heading(doc, "3.4 Model Training Module", 2)
    add_long_paragraphs(doc, [
        "The Model Training Module compiles and trains each candidate using TensorFlow/Keras. Classification models use Adam optimizer, categorical cross-entropy loss, and accuracy. Regression models use Adam optimizer, mean squared error loss, and MSE.",
        "Training data is supplied through tf.data pipelines, and an early stopping callback limits unnecessary computation. A custom stream callback forwards progress updates to the event queue for real-time interface updates."
    ])
    add_heading(doc, "3.5 Model Evaluation Module", 2)
    add_long_paragraphs(doc, [
        "The Model Evaluation Module evaluates trained candidates on validation and test sets after training. It records loss and the task-specific metric returned by the compiled model.",
        "For classification, the preserved values include accuracy and loss. For regression, the preserved values include MSE and loss. These results are later used for ranking, reporting, and model comparison."
    ])
    add_heading(doc, "3.6 Model Selection Module", 2)
    add_long_paragraphs(doc, [
        "The Model Selection Module chooses the best candidate using validation performance rather than final test output. For classification, higher validation accuracy is preferred. For regression, lower validation MSE is preferred through the implemented ranking rule.",
        "This design keeps selection logic consistent with accepted ML practice and ensures that the saved best model is linked to the preserved candidate results and readable summary."
    ])
    add_heading(doc, "3.7 Model Storage Module", 2)
    add_long_paragraphs(doc, [
        "The Model Storage Module saves each trained candidate as a `.keras` file. It also saves preprocessing and metadata artifacts as a `.joblib` bundle and writes run-level summary information to `training_report.json`.",
        "The saved joblib artifact includes the fitted preprocessor, feature columns, feature schema, task type, target name, and label encoder when classification is used. These files support downloads and later prediction reuse."
    ])
    add_heading(doc, "3.8 Prediction Module", 2)
    add_long_paragraphs(doc, [
        "The Prediction Module loads saved artifacts, rebuilds a single-row input DataFrame using saved feature ordering and schema information, applies the stored preprocessor, and runs inference using the cached best model.",
        "For classification, the module returns the predicted label and class probabilities. For regression, it returns a numeric prediction. The module accepts JSON or form-data input according to the implemented route logic."
    ])
    add_heading(doc, "3.9 Frontend UI Module", 2)
    add_long_paragraphs(doc, [
        "The Frontend UI Module consists of the training page and the prediction page. The training page supports dataset upload, NAS mode selection, and live progress display. The prediction page retrieves feature schema metadata and builds the form dynamically.",
        "This module depends on the backend for model readiness, progress updates, model comparison data, and prediction responses, but keeps user interaction separate from training logic."
    ])

    add_heading(doc, "4. Data Flow", 1)
    add_numbered(doc, [
        "Dataset is uploaded through the training interface using CSV or Excel format.",
        "The backend validates file type, stores the file, and starts background training.",
        "The Data Pipeline Module reads the dataset, handles missing values, encodes categorical fields, normalizes numeric fields, and detects classification or regression.",
        "The processed dataset is split into train, validation, and test partitions using an effective 60/20/20 workflow.",
        "The NAS Engine generates dense neural network candidates according to the selected search mode.",
        "The Model Training Module compiles and trains candidates using Adam optimizer and task-appropriate loss/metric configuration.",
        "The Model Evaluation Module records validation and test outcomes for every candidate.",
        "The Model Selection Module selects the best candidate using validation performance.",
        "The Model Storage Module saves candidate models, best pipeline artifacts, and the training report.",
        "The Prediction Module and prediction page reuse the saved feature schema, preprocessing bundle, and best model for inference.",
    ])

    add_heading(doc, "5. API Design", 1)
    add_heading(doc, "5.1 `/upload`", 2)
    add_long_paragraphs(doc, [
        "Method: POST. This endpoint accepts the dataset file and training configuration fields. It validates the input file type, validates NAS type, clamps candidate count, batch size, and epoch values, and starts the background training thread.",
        "The endpoint returns a JSON response confirming that training has started or an error response when the file is missing, the format is unsupported, the NAS type is invalid, or another run is already active."
    ])
    add_heading(doc, "5.2 `/stream`", 2)
    add_long_paragraphs(doc, [
        "Method: GET. This endpoint provides Server-Sent Events for real-time training updates.",
        "The endpoint streams status messages, model information, epoch-end results, and final candidate result events from the event queue to the frontend."
    ])
    add_heading(doc, "5.3 `/api/model-comparison`", 2)
    add_long_paragraphs(doc, [
        "Method: GET. This endpoint returns the saved model comparison information required by the frontend after training.",
        "The response includes task type, best model information, all saved model summaries, and readable summary text derived from the saved report."
    ])
    add_heading(doc, "5.4 `/predict`", 2)
    add_long_paragraphs(doc, [
        "Method: POST. This endpoint accepts prediction input values, reconstructs the row according to the saved feature schema, applies the saved preprocessor, and runs inference using the cached best model.",
        "The response returns a predicted label and probabilities for classification or a numeric prediction for regression. If no trained artifacts are available, the endpoint returns an error response."
    ])

    add_heading(doc, "6. Design Constraints", 1)
    add_bullets(doc, [
        "Only dense neural networks are implemented.",
        "The search space is intentionally limited to 1–4 hidden layers.",
        "Hidden-layer units are limited to 16, 32, 64, and 128.",
        "Hidden activations are limited to ReLU and Tanh.",
        "The implemented optimizer is Adam with learning rate 0.001.",
        "The system does not implement CNN, RNN, GAN, or other advanced architectures.",
        "The system is designed for tabular data input from CSV and Excel files only.",
    ])

    add_heading(doc, "7. Future Enhancements", 1)
    add_bullets(doc, [
        "Expose more of the already implemented safe configuration controls in the UI.",
        "Add richer artifact browsing and run-history presentation for repeated academic use.",
        "Improve generated documentation and report views from the saved training report.",
        "Strengthen automated validation and testing around upload, training, and prediction workflows.",
        "Enhance interface guidance around feature schema and prediction input entry.",
    ])

    output = DOC_DIR / "3_SDD_DOCUMENT.docx"
    output.unlink(missing_ok=True)
    doc.save(output)


def generate_diary() -> None:
    doc = new_document("Project Diary", f"{PROJECT_TITLE} Development Record")
    weeks = [
        ("Week 1", "Requirement study and workspace familiarization"),
        ("Week 2", "Dataset handling design"),
        ("Week 3", "NAS search-space study"),
        ("Week 4", "Backend route analysis"),
        ("Week 5", "Frontend and UX review"),
        ("Week 6", "Artifact validation"),
        ("Week 7", "Literature consolidation"),
        ("Week 8", "Report correction"),
        ("Week 9", "Specification correction"),
        ("Week 10", "Diary and mapping correction"),
        ("Week 11", "Validation and consistency check"),
        ("Week 12", "Submission preparation"),
    ]
    for label, title in weeks:
        add_heading(doc, f"{label}: {title}", 1)
        add_paragraph(doc, "Estimated Effort: 8-12 hours")
        add_heading(doc, "Work Done", 2)
        add_bullets(doc, [
            f"Reviewed the project area related to {title.lower()}.",
            "Verified that documentation remains grounded in the real implementation.",
        ])
        add_heading(doc, "Challenges Faced", 2)
        add_bullets(doc, [
            "Maintaining consistency between generated documents and evolving project artifacts.",
        ])
        add_heading(doc, "Learning Outcomes", 2)
        add_bullets(doc, [
            "Improved understanding of full-stack ML workflow design.",
            "Recognized the importance of evidence-based academic reporting.",
        ])
        add_heading(doc, "Next Plan", 2)
        add_paragraph(doc, "Proceed to the next project documentation and validation stage.")
    output = DOC_DIR / "4_PROJECT_DIARY.docx"
    output.unlink(missing_ok=True)
    doc.save(output)


def generate_mapping() -> None:
    doc = new_document("PO–PSO–SDG Mapping Document", PROJECT_TITLE)
    add_heading(doc, "1. Project Summary", 1)
    add_paragraph(doc, "The project integrates web development, preprocessing, model search, evaluation, and prediction support for tabular medical datasets.")
    add_heading(doc, "2. Program Outcomes (PO)", 1)
    add_report_table(doc, ["PO", "Outcome", "Level", "Evidence"], [
        ["PO1", "Engineering knowledge", "3", "Applied ML, preprocessing, and web integration"],
        ["PO2", "Problem analysis", "3", "Analysed dataset, task type, and model behavior"],
        ["PO3", "Design and development", "3", "Built layered application and training workflow"],
        ["PO4", "Investigation", "3", "Compared candidate models and studied literature"],
        ["PO5", "Modern tool usage", "3", "Used Flask, TensorFlow, scikit-learn, Pandas"],
        ["PO10", "Communication", "3", "Prepared report, SRS, SDD, diary, and mapping documents"],
    ])
    add_heading(doc, "3. Program Specific Outcomes (PSO)", 1)
    add_report_table(doc, ["PSO", "Statement", "Level"], [
        ["PSO1", "Design and implement an AI/ML system for tabular prediction", "3"],
        ["PSO2", "Build an end-to-end application with data, model, and UI integration", "3"],
    ])
    add_heading(doc, "4. SDG Alignment", 1)
    add_bullets(doc, [
        "SDG 3: Good Health and Well-Being",
        "SDG 4: Quality Education",
        "SDG 9: Industry, Innovation and Infrastructure",
    ])
    output = DOC_DIR / "5_PO_PSO_SDG_MAPPING.docx"
    output.unlink(missing_ok=True)
    doc.save(output)


def main() -> None:
    generate_project_report()
    generate_srs()
    generate_sdd()
    generate_diary()
    generate_mapping()
    print("Regenerated academic documents successfully.")


if __name__ == "__main__":
    main()
